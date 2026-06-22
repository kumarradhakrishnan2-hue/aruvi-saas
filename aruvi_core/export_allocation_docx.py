"""
Period Allocation Report — DOCX export using python-docx.

Pure-Python (no Node, no headless browser) so it runs on a hosted cloud API and
produces a fully editable, email-shareable Word document. Mirrors the PDF layout
exactly (same three sections, same wording, no dark bands):

  Header: "Aruvi.LESSON STUDIO" + "NCF 2023 aligned" (left), "Allocation &
          Competency report" + Grade·Subject·Date (right), thin rule below.
  Summary strip: Chapters / Periods / Total time / Period types.
  Executive summary (lowercase serif heading).
  Allocation details — the periods table (#, Chapter, per-duration, Total, basis).
  Competency report — per-chapter blocks with competency tables.
"""

from __future__ import annotations

from io import BytesIO

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from .report_competency import (
    CompetencyAllocationReport, grade_roman, subject_display, date_long,
    executive_summary_paragraphs,
)

INK = RGBColor(0x1A, 0x19, 0x17)
PINE = RGBColor(0x16, 0x44, 0x36)
CLAY = RGBColor(0xB6, 0x5A, 0x31)
GRAY = RGBColor(0x88, 0x88, 0x88)
MUTED = RGBColor(0xAA, 0xAA, 0xAA)
BODY = RGBColor(0x33, 0x33, 0x33)
SERIF = "Georgia"
ALT_HEX = "FAF9F7"
TOTAL_HEX = "F4F2EE"


def _bg(cell, hex_color):
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear"); shd.set(qn("w:fill"), hex_color)
    cell._tc.get_or_add_tcPr().append(shd)


def _no_borders(table):
    """Remove all cell borders, then we add only the hairlines we want."""
    tblPr = table._tbl.tblPr
    b = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        e = OxmlElement(f"w:{edge}")
        e.set(qn("w:val"), "nil")
        b.append(e)
    tblPr.append(b)


def _hairlines(table, color="DDDDDD"):
    """Light horizontal hairlines only (no vertical lines, no dark grid)."""
    tblPr = table._tbl.tblPr
    b = OxmlElement("w:tblBorders")
    for edge in ("top", "bottom", "insideH"):
        e = OxmlElement(f"w:{edge}")
        e.set(qn("w:val"), "single"); e.set(qn("w:sz"), "4")
        e.set(qn("w:space"), "0"); e.set(qn("w:color"), color)
        b.append(e)
    for edge in ("left", "right", "insideV"):
        e = OxmlElement(f"w:{edge}")
        e.set(qn("w:val"), "nil")
        b.append(e)
    tblPr.append(b)


def _set_widths(table, widths_in):
    """Reliably fix column widths in Word: disable autofit, set a fixed layout, and
    write the width on EVERY cell (Word ignores table.columns[i].width on its own)."""
    table.autofit = False
    table.allow_autofit = False
    # fixed layout so Word honours the per-cell widths
    tblPr = table._tbl.tblPr
    layout = OxmlElement("w:tblLayout"); layout.set(qn("w:type"), "fixed")
    tblPr.append(layout)
    for row in table.rows:
        for i, w in enumerate(widths_in):
            if i < len(row.cells):
                row.cells[i].width = Inches(w)


def _run(p, text, *, size=9, bold=False, color=INK, italic=False, font="Calibri", spacing=None):
    r = p.add_run(text)
    r.font.size = Pt(size); r.font.bold = bold; r.font.italic = italic
    r.font.color.rgb = color; r.font.name = font
    return r


def _cell(cell, text, *, size=8, bold=False, color=INK, align="left", font="Calibri"):
    cell.paragraphs[0].clear()
    p = cell.paragraphs[0]
    p.alignment = {"left": WD_ALIGN_PARAGRAPH.LEFT, "right": WD_ALIGN_PARAGRAPH.RIGHT,
                   "center": WD_ALIGN_PARAGRAPH.CENTER}[align]
    p.paragraph_format.space_before = Pt(1); p.paragraph_format.space_after = Pt(1)
    _run(p, str(text), size=size, bold=bold, color=color, font=font)


def _section_head(doc, text):
    """(point 6) Lowercase serif heading, two notches above chapter title size, with a rule."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12); p.paragraph_format.space_after = Pt(5)
    _run(p, text, size=14, bold=True, color=INK, font=SERIF)
    _rule(p, color="1A1917", sz=10)  # one notch thicker than the chapter rule


def _g(v) -> str:
    try:
        f = float(v)
        return str(int(f)) if f == int(f) else str(f)
    except (TypeError, ValueError):
        return str(v)


def export_allocation_report_docx(report: CompetencyAllocationReport) -> bytes:
    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = Inches(0.7); sec.bottom_margin = Inches(0.7)
    sec.left_margin = Inches(0.7); sec.right_margin = Inches(0.7)

    g = grade_roman(report.grade)
    subj = subject_display(report.subject)

    # ── Header (point 5: site logo) + (point 4: renamed) ──
    ht = doc.add_table(rows=1, cols=2)
    _no_borders(ht)
    ht.columns[0].width = Inches(4.0); ht.columns[1].width = Inches(3.1)
    lc = ht.rows[0].cells[0]
    lc.paragraphs[0].clear()
    lp = lc.paragraphs[0]; lp.paragraph_format.space_after = Pt(0)
    _run(lp, "Aruvi", size=17, bold=True, color=PINE, font=SERIF)
    _run(lp, ".", size=17, bold=True, italic=True, color=CLAY, font=SERIF)
    _run(lp, "  LESSON STUDIO", size=7, color=GRAY, font="Calibri")
    lp2 = lc.add_paragraph(); lp2.paragraph_format.space_before = Pt(1)
    _run(lp2, "NCF 2023 aligned", size=8, italic=True, color=GRAY, font=SERIF)
    rc = ht.rows[0].cells[1]
    rc.paragraphs[0].clear()
    rp = rc.paragraphs[0]; rp.alignment = WD_ALIGN_PARAGRAPH.RIGHT; rp.paragraph_format.space_after = Pt(0)
    _run(rp, "Allocation & Competency report", size=12, bold=True, color=PINE, font=SERIF)
    rp2 = rc.add_paragraph(); rp2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    _run(rp2, f"Grade {g} · {subj} · {date_long(report.generated_at)}", size=8, color=GRAY)

    # (point 2) gap below logo, then a thin rule
    gap = doc.add_paragraph(); gap.paragraph_format.space_after = Pt(2)
    _rule(gap, color="1A1917", sz=16)  # heavy 2px prototype-style rule under the logo

    # ── Summary strip (point 1) ──
    total_mins = report.total_minutes
    hrs, rem = divmod(total_mins, 60)
    time_str = f"{hrs}h {rem}min" if rem else f"{hrs}h"
    ptype = " · ".join(f"{t.count}×{t.minutes}min" for t in report.sorted_types) or "—"
    stats = [(str(len(report.chapters)), "Chapters"), (str(report.total_periods), "Periods"),
             (time_str, "Total time"), (ptype, "Period types")]
    stt = doc.add_table(rows=2, cols=len(stats)); stt.alignment = WD_TABLE_ALIGNMENT.CENTER
    _hairlines(stt, "DDDDDD")
    for i, (v, k) in enumerate(stats):
        _cell(stt.rows[0].cells[i], v, size=12, bold=True, align="center")
        _cell(stt.rows[1].cells[i], k.upper(), size=6, color=GRAY, align="center")
    doc.add_paragraph().paragraph_format.space_after = Pt(2)

    # ── Executive summary (point 7) ──
    _section_head(doc, "Executive summary")
    for para in executive_summary_paragraphs(report):
        pp = doc.add_paragraph(); pp.paragraph_format.space_after = Pt(5)
        _run(pp, para, size=9.5, color=BODY)

    # ── Allocation details (point 9) ──
    _section_head(doc, "Allocation details")
    _allocation_table(doc, report)

    # ── Competency report (point 10) ──
    _section_head(doc, "Competency report")
    for ch in report.chapters:
        _chapter_block(doc, ch, report)

    buf = BytesIO(); doc.save(buf); buf.seek(0)
    return buf.getvalue()


def _allocation_table(doc, report):
    """(point 9) Clean periods table — no dark bands; light striping + a Total row."""
    types = report.sorted_types
    metric_head = "Effort Index" if report.is_effort else "Competency Weight"
    headers = ["#", "Chapter"] + [f"{t.minutes}-min Periods" for t in types] + ["Total Periods", metric_head]
    t = doc.add_table(rows=1, cols=len(headers)); t.alignment = WD_TABLE_ALIGNMENT.CENTER
    _hairlines(t, "DDDDDD")
    for i, h in enumerate(headers):
        _cell(t.rows[0].cells[i], h, size=6.5, color=GRAY,
              align="left" if i == 1 else "center")

    tot_by_dur = {tp.minutes: 0 for tp in types}
    tot_periods = 0
    for idx, ch in enumerate(report.chapters, 1):
        cells = t.add_row().cells
        c = 0
        _cell(cells[c], idx, size=7, color=MUTED, align="center"); c += 1
        _cell(cells[c], ch.chapter_title, size=8, align="left"); c += 1
        for tp in types:
            v = ch.periods_by_duration.get(tp.minutes, 0)
            tot_by_dur[tp.minutes] += v
            _cell(cells[c], v, size=8, align="center"); c += 1
        tot_periods += ch.total_periods
        _cell(cells[c], ch.total_periods, size=8, bold=True, align="center"); c += 1
        metric = ch.effort_index if report.is_effort else ch.chapter_weight
        _cell(cells[c], "" if metric in (None, "") else _g(metric), size=8, bold=True, align="center")
        if idx % 2 == 0:
            for cell in cells:
                _bg(cell, ALT_HEX)
    # Total row
    cells = t.add_row().cells
    c = 0
    _cell(cells[c], "", size=8); c += 1
    _cell(cells[c], "Total", size=8, bold=True, align="left"); c += 1
    for tp in types:
        _cell(cells[c], tot_by_dur[tp.minutes], size=8, bold=True, align="center"); c += 1
    _cell(cells[c], tot_periods, size=8, bold=True, align="center"); c += 1
    _cell(cells[c], "", size=8)
    for cell in cells:
        _bg(cell, TOTAL_HEX)
    # Fixed widths (≈7.1in total): #, Chapter, each duration, Total Periods, metric.
    n_dur = len(types)
    dur_w = 1.0
    chapter_w = max(1.6, 7.1 - 0.4 - dur_w * n_dur - 1.1 - 1.2)
    _set_widths(t, [0.4, chapter_w] + [dur_w] * n_dur + [1.1, 1.2])


def _chapter_block(doc, ch, report):
    period_cells = " · ".join(
        f"{ch.periods_by_duration.get(t.minutes, 0)}×{t.minutes}min" for t in report.sorted_types
    ) or "—"
    # Chapter header (point 8: Ch NN dark like title; point 3: allocation plain, no dark band)
    hp = doc.add_paragraph(); hp.paragraph_format.space_before = Pt(10); hp.paragraph_format.space_after = Pt(2)
    _run(hp, f"Ch {str(ch.chapter_number).zfill(2)}  ", size=10, bold=True, color=INK, font=SERIF)
    _run(hp, ch.chapter_title, size=10, bold=True, color=INK, font=SERIF)
    _run(hp, f"    {period_cells} · {ch.total_periods} periods · {ch.total_minutes}min",
         size=7.5, color=GRAY)
    if not report.is_effort and ch.chapter_weight not in (None, ""):
        _run(hp, f"   Weight {ch.chapter_weight}", size=7.5, color=MUTED)
    _rule(hp, color="1A1917", sz=6)  # thin chapter rule (thinner than the section-head rule)

    show_wt = (not report.is_effort) and any(c.weight is not None for c in ch.competencies)
    if not ch.competencies:
        np = doc.add_paragraph(); _run(np, "No competency entries for this chapter.", size=8, color=MUTED, italic=True)
        return

    ncols = 5 if show_wt else 4
    t = doc.add_table(rows=1, cols=ncols); t.alignment = WD_TABLE_ALIGNMENT.LEFT
    _hairlines(t, "E6E6E6")
    headers = ["#", "Code", "Competency", "Justification"] + (["Weight"] if show_wt else [])
    # (point 3) header is a LIGHT row with grey text + underline, NOT a dark band
    for i, h in enumerate(headers):
        _cell(t.rows[0].cells[i], h, size=6, color=GRAY,
              align="center" if i in (0, 4) else "left")
    for idx, c in enumerate(ch.competencies, 1):
        cells = t.add_row().cells
        _cell(cells[0], idx, size=7, color=MUTED, align="center")
        _cell(cells[1], c.c_code, size=8, bold=True)
        _cell(cells[2], c.description, size=8)
        _cell(cells[3], c.justification, size=8)
        if show_wt:
            w = int(c.weight or 0)
            _cell(cells[4], "●" * w + "○" * (3 - w), size=8, align="center")
    # Page content width ≈ 7.1in (Letter − 0.7in margins). Match the PDF proportions so
    # the # and Code columns stay narrow and Competency/Justification get the room.
    if show_wt:
        _set_widths(t, [0.4, 0.8, 2.1, 3.0, 0.8])   # #, Code, Competency, Justification, Weight
    else:
        _set_widths(t, [0.4, 0.85, 2.3, 3.55])       # #, Code, Competency, Justification


def _rule(p, *, color="1A1917", sz=6):
    """Add a bottom border to an existing paragraph `p`."""
    pPr = p._p.get_or_add_pPr()
    pbdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single"); bottom.set(qn("w:sz"), str(sz))
    bottom.set(qn("w:space"), "2"); bottom.set(qn("w:color"), color)
    pbdr.append(bottom); pPr.append(pbdr)
