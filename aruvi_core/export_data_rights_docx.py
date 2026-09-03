"""
Data-rights export — the teacher's own data as ONE editable Word document.

Administrative architecture Step 4, reformatted 2026-08-22 per founder review of the
first live export. Visual language is the ALLOCATION REPORT's, reused directly (its
helpers are imported, not copied, so the two documents cannot drift): Aruvi.LESSON
STUDIO header table, heavy rule, serif section heads with hairline rules, bordered-
hairline tables with fixed widths, Georgia/Calibri, pine/ink/gray.

Layout (each item is a founder direction, 2026-08-22):
  1. House format throughout.
  2. A PURPOSE statement at the top — the spirit of "this is yours", no regulation
     names.
  3. Teaching profile as a TABLE: Subject · Class · Sections.
  4. Chapter notes NUMBERED serially, each a separated block with a bold identity
     line (Subject · Class · Chapter no. — Chapter name) above the note text.
  5. Teaching state as a TABLE: No. · Subject · Class · Chapter · Sections · Status —
     no filenames, no canonical identities, no period internals.
  6. The closing "About lesson plan content" note also advises exporting BEFORE an
     account deletion, since Aruvi then removes all personal data within a short time.

The builder consumes the plain-dict payload assembled by
adapters/data_rights_service_file.py — it never touches storage itself, so the
partner's cloud adapter can reuse it unchanged.
"""

from __future__ import annotations

from datetime import datetime, timezone
from io import BytesIO
from typing import Any, Dict, List

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from . import brand   # the MEYY wordmark raster (2026-09-03)

from .export_allocation_docx import (
    _cell, _hairlines, _no_borders, _rule, _run, _section_head, _set_widths,
    INK, PINE, CLAY, GRAY, BODY, SERIF, ALT_HEX, _bg,
)
from .report_competency import date_long, subject_display


def _class_of(grade_slug: str) -> str:
    """Grade slug → the teacher-facing class label ("iv" → "IV")."""
    return str(grade_slug or "").replace("grade", "").replace("_", " ").strip().upper()


def _subj(slug: str) -> str:
    return subject_display(slug) if slug else ""


def _body_para(doc, text, *, size=9.5, color=BODY, italic=False, space_after=5):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    _run(p, text, size=size, color=color, italic=italic, font=SERIF)
    return p


def _striped(table):
    """Light alternating row fill under hairlines — the allocation table's look."""
    _hairlines(table, "DDDDDD")
    for i, row in enumerate(table.rows[1:], start=1):
        if i % 2 == 0:
            for c in row.cells:
                _bg(c, ALT_HEX)


def build_export_docx(payload: Dict[str, Any]) -> bytes:
    """payload (assembled by the data-rights service):
    {account, exported_at, profile: {subjects:[...]},
     years: [{year_id,
              notes: {key: {text, updated_at, chapter_title}},
              teaching: [{subject, grade, chapter_number, chapter_title,
                          sections: [{tag, status}]}], …}]}
    """
    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = Inches(0.7); sec.bottom_margin = Inches(0.7)
    sec.left_margin = Inches(0.7); sec.right_margin = Inches(0.7)

    acct = payload.get("account") or {}
    try:
        when = datetime.fromisoformat(str(payload.get("exported_at", "")))
    except ValueError:
        when = datetime.now(timezone.utc)

    # ── Header — the allocation report's, retitled ──
    ht = doc.add_table(rows=1, cols=2)
    _no_borders(ht)
    ht.columns[0].width = Inches(4.0); ht.columns[1].width = Inches(3.1)
    lc = ht.rows[0].cells[0]
    lc.paragraphs[0].clear()
    lp = lc.paragraphs[0]; lp.paragraph_format.space_after = Pt(0)
    brand.add_wordmark(lp, 17)   # the MEYY wordmark (2026-09-03)
    lp2 = lc.add_paragraph(); lp2.paragraph_format.space_before = Pt(1)   # kicker beneath the mark; NCF line gone (founder, 2026-09-03)
    _run(lp2, "LESSON STUDIO", size=7, color=GRAY, font="Calibri")
    rc = ht.rows[0].cells[1]
    rc.paragraphs[0].clear()
    rp = rc.paragraphs[0]; rp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    rp.paragraph_format.space_after = Pt(0)
    _run(rp, "Your data export", size=12, bold=True, color=PINE, font=SERIF)
    rp2 = rc.add_paragraph(); rp2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    _run(rp2, f"{acct.get('display_name') or acct.get('account_id') or ''}"
              f" · {date_long(when)}", size=8, color=GRAY)
    gap = doc.add_paragraph(); gap.paragraph_format.space_after = Pt(2)
    _rule(gap, color="1A1917", sz=16)

    # ── Purpose (founder point 2 — the spirit, no regulation names) ──
    _body_para(doc,
        "Everything you create in Meyy belongs to you — your notes, your teaching "
        "profile, your progress through the year. This document is a complete, "
        "editable copy of all of it, in one place, so your work is always yours to "
        "keep, carry and continue — with Meyy or without it.",
        italic=True, space_after=10)

    # ── Your account ──
    _section_head(doc, "Your account")
    rows = [(label, acct.get(key)) for label, key in (
        ("Name", "display_name"), ("User ID", "account_id"), ("Email", "email"),
        ("Phone", "phone"), ("School", "school_name"), ("Member since", "created_at"))
        if acct.get(key)]
    # The agreement she accepted (2026-08-27). Part of what Aruvi holds about her, so it
    # belongs in what she can download — and it is the one item on this page that is
    # deliberately RETAINED after an erase, which the erasure receipt says in the same
    # words. Read off the account's `consent` mirror; the ledger is the authority.
    _consent = acct.get("consent") or {}
    if _consent.get("accepted_at"):
        rows.append(("Agreement accepted",
                     f"User Agreement v{_consent.get('policy_version') or '—'} "
                     f"on {str(_consent.get('accepted_at'))[:10]}"))
    if rows:
        at = doc.add_table(rows=len(rows), cols=2)
        _hairlines(at, "DDDDDD")
        _set_widths(at, [1.6, 5.5])
        for i, (label, value) in enumerate(rows):
            v = str(value)[:10] if label == "Member since" else str(value)
            _cell(at.rows[i].cells[0], label, size=9, color=GRAY)
            _cell(at.rows[i].cells[1], v, size=9)
    else:
        _body_para(doc, "No account record found.", color=GRAY, italic=True)

    # ── Your teaching profile — table (founder point 3) ──
    _section_head(doc, "Your teaching profile")
    prows: List[List[str]] = []
    for s in ((payload.get("profile") or {}).get("subjects")) or []:
        for g in s.get("grades") or []:
            secs = ", ".join(x.get("tag", "") for x in (g.get("sections") or [])) or "—"
            prows.append([s.get("name", ""), _class_of(g.get("grade", "")), secs])
    if prows:
        pt = doc.add_table(rows=1 + len(prows), cols=3)
        _striped_headers(pt, ["Subject", "Class", "Sections"])
        _set_widths(pt, [3.0, 1.2, 2.9])
        for i, r in enumerate(prows, start=1):
            for j, v in enumerate(r):
                _cell(pt.rows[i].cells[j], v, size=9)
        _striped(pt)
    else:
        _body_para(doc, "No teaching profile on record.", color=GRAY, italic=True)

    # ── Per academic year ──
    for yr in payload.get("years") or []:
        _section_head(doc, f"Academic year {yr.get('year_id', '')}")

        # Chapter notes — numbered, separated blocks (founder point 4).
        _sub_head(doc, "Your chapter notes")
        notes = yr.get("notes") or {}
        if notes:
            for n_i, key in enumerate(sorted(notes), start=1):
                n = notes[key]
                subj, grade, chapter = (key.split("/") + ["", "", ""])[:3]
                title = n.get("chapter_title") or ""
                head = (f"{n_i}.  {_subj(subj)} · Class {_class_of(grade)} · "
                        f"Chapter {chapter}" + (f" — {title}" if title else ""))
                hp = doc.add_paragraph()
                hp.paragraph_format.space_before = Pt(8)
                hp.paragraph_format.space_after = Pt(2)
                _run(hp, head, size=10, bold=True, color=INK, font=SERIF)
                _body_para(doc, n.get("text", ""), space_after=2)
                if n.get("updated_at"):
                    _body_para(doc, f"Last edited {str(n['updated_at'])[:10]}",
                               size=8, color=GRAY, italic=True, space_after=8)
        else:
            _body_para(doc, "No notes this year.", color=GRAY, italic=True)

        # Teaching state — table (founder point 5: no filenames, no periods).
        _sub_head(doc, "Your teaching state")
        teaching = yr.get("teaching") or []
        if teaching:
            tt = doc.add_table(rows=1 + len(teaching), cols=6)
            _striped_headers(tt, ["No.", "Subject", "Class", "Chapter", "Sections",
                                  "Status"])
            _set_widths(tt, [0.4, 1.7, 0.6, 2.0, 0.9, 1.5])
            for i, row in enumerate(teaching, start=1):
                num = row.get("chapter_number")
                title = row.get("chapter_title") or ""
                chap = (f"Ch. {num}" if num else "—") + (f" — {title}" if title else "")
                secs = ", ".join(s["tag"] for s in row.get("sections") or [])
                status = " · ".join(f"{s['tag']}: {s['status']}"
                                    for s in row.get("sections") or [])
                for j, v in enumerate([str(i), _subj(row.get("subject", "")),
                                       _class_of(row.get("grade", "")), chap, secs,
                                       status]):
                    _cell(tt.rows[i].cells[j], v, size=8.5)
            _striped(tt)
        else:
            _body_para(doc, "No teaching activity this year.", color=GRAY, italic=True)

    # ── Messages you sent us (2026-08-27) ──
    # Only when there are some: an empty "Messages you sent us" heading on every export
    # would tell most teachers they were supposed to have written to somebody. Placed
    # AFTER the years, because it is the least of what she came for.
    support = payload.get("support") or []
    if support:
        _section_head(doc, "Messages you sent us")
        _body_para(doc,
            "Support messages you have written to Meyy, newest first, with the "
            "reference each was given.", size=9, color=GRAY, italic=True, space_after=6)
        for s in support:
            hp = doc.add_paragraph()
            hp.paragraph_format.space_before = Pt(8)
            hp.paragraph_format.space_after = Pt(2)
            label = s.get("category_label") or s.get("category") or ""
            _run(hp, f"{s.get('reference', '')}  ·  {label}",
                 size=10, bold=True, color=INK, font=SERIF)
            _body_para(doc, s.get("message", ""), space_after=2)
            if s.get("created_at"):
                _body_para(doc, f"Sent {str(s['created_at'])[:10]}",
                           size=8, color=GRAY, italic=True, space_after=8)

    # ── Closing note (founder point 6: current text + pre-deletion advisory) ──
    _section_head(doc, "About lesson plan content")
    _body_para(doc,
        "Lesson plans and assessments themselves are Meyy's shared library content "
        "and are not personal data; export them any time as PDFs from the app. "
        "If you are considering deleting your Meyy account, we suggest you export "
        "this document first and keep it safely: when an account is deleted, Meyy "
        "removes all personal data relating to your activity within a short period "
        "of time, and it cannot be recovered afterwards.",
        size=9, color=GRAY, italic=True)

    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _sub_head(doc, text):
    """Quieter sub-heading under a year's section head."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8); p.paragraph_format.space_after = Pt(4)
    _run(p, text, size=11, bold=True, color=PINE, font=SERIF)


def _striped_headers(table, headers):
    """Header row in the allocation table's voice: small caps-ish gray labels."""
    for j, h in enumerate(headers):
        _cell(table.rows[0].cells[j], h.upper(), size=7, bold=True, color=GRAY)
