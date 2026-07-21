"""
Lesson Plan / Assessment / Integrated — DOCX (Word) export via python-docx.

The editable-Word sibling of the three PDF exporters. Pure-Python (no Node, no
headless browser) so it runs on a hosted cloud API and produces a fully editable,
email-shareable Word document. Content, answer-layer, and unit-scoping parity with
the PDFs; styling is structural (headings, shaded bands, brand colours) rather than
pixel-perfect, since Word is editable.

Public entry points mirror the PDF exporters:
  export_lesson_plan_docx(view, *, competencies, plan_date, generated_at)
  export_assessment_docx(view, *, include_answers, assessment_type, plan_date)
  export_integrated_docx(view, *, include_answers, unit_number, competencies, plan_date)

`view` is ViewModel.to_dict(); the same enrichment (competencies, plan_date) the
API assembles for the PDFs is passed straight through.
"""

from __future__ import annotations

from datetime import datetime
from io import BytesIO
from typing import Any, Dict, List, Optional

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from .report_competency import grade_roman, subject_display
from .export_lesson_pdf import (
    targeted_competencies, _duration_breakdown, _group_word as _lp_group_word,
    _phase_duration, AXIS_INFO,
)
from .export_assessment_pdf import intro_paragraph, CTX_WORD, question_type_full
from .export_integrated_pdf import _items_by_anchor, _iter_group_periods

# ── palette ──────────────────────────────────────────────────────────────────
INK = RGBColor(0x1A, 0x19, 0x17)
PINE = RGBColor(0x16, 0x44, 0x36)
CLAY = RGBColor(0xB6, 0x5A, 0x31)
SOFT = RGBColor(0x6B, 0x6A, 0x63)
BODY = RGBColor(0x2A, 0x2A, 0x2A)
G_ACCENT = RGBColor(0x0F, 0x6E, 0x56)
G_DARK = RGBColor(0x0C, 0x3A, 0x2E)
KRAFT_HEX = "EBE3D3"
GHEAD_HEX = "DCEAE3"
GTINT_HEX = "EEF6F1"
LINE_HEX = "DDDDDD"
SERIF = "Georgia"
SANS = "Arial"


# ── low-level helpers ────────────────────────────────────────────────────────

def _run(p, text, *, bold=False, italic=False, size=9.5, color=INK, font=SANS, caps=False):
    r = p.add_run(str(text).upper() if caps else str(text))
    r.bold = bold
    r.italic = italic
    r.font.size = Pt(size)
    r.font.color.rgb = color
    r.font.name = font
    return r


def _para(doc_or_cell, *, space_before=0, space_after=2, align=None):
    p = doc_or_cell.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    if align is not None:
        p.alignment = align
    return p


def _bg(cell, hex_color):
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), hex_color)
    cell._tc.get_or_add_tcPr().append(shd)


def _no_borders(table):
    b = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        e = OxmlElement(f"w:{edge}")
        e.set(qn("w:val"), "nil")
        b.append(e)
    table._tbl.tblPr.append(b)


def _fixed_table(table, col_inches):
    """Force real fixed column widths (Word/LibreOffice ignore per-cell width unless the table
    layout is 'fixed' AND the grid columns are set). col_inches: list of widths in inches."""
    table.autofit = False
    tbl = table._tbl
    tblPr = tbl.tblPr
    for el in tblPr.findall(qn("w:tblLayout")):
        tblPr.remove(el)
    layout = OxmlElement("w:tblLayout"); layout.set(qn("w:type"), "fixed"); tblPr.append(layout)
    for el in tblPr.findall(qn("w:tblW")):
        tblPr.remove(el)
    tblW = OxmlElement("w:tblW")
    tblW.set(qn("w:w"), str(int(sum(col_inches) * 1440))); tblW.set(qn("w:type"), "dxa")
    tblPr.append(tblW)
    grid = tbl.find(qn("w:tblGrid"))
    if grid is not None:
        for gc in grid.findall(qn("w:gridCol")):
            grid.remove(gc)
        for w in col_inches:
            gc = OxmlElement("w:gridCol"); gc.set(qn("w:w"), str(int(w * 1440))); grid.append(gc)
    for row in table.rows:
        for i, w in enumerate(col_inches):
            row.cells[i].width = Inches(w)


def _hairlines(table, color=LINE_HEX):
    b = OxmlElement("w:tblBorders")
    for edge in ("top", "bottom", "insideH"):
        e = OxmlElement(f"w:{edge}")
        e.set(qn("w:val"), "single"); e.set(qn("w:sz"), "4")
        e.set(qn("w:space"), "0"); e.set(qn("w:color"), color)
        b.append(e)
    for edge in ("left", "right", "insideV"):
        e = OxmlElement(f"w:{edge}")
        e.set(qn("w:val"), "nil")
        b.append(e)
    table._tbl.tblPr.append(b)


def _shaded_band(doc, hex_color, kicker, kicker_color, title, title_color, *, title_size=13):
    """A full-width one-cell shaded band with a small kicker over a title (stage bands)."""
    t = doc.add_table(rows=1, cols=1)
    _no_borders(t)
    cell = t.cell(0, 0)
    _bg(cell, hex_color)
    cell.paragraphs[0].paragraph_format.space_after = Pt(0)
    if kicker:
        kp = cell.paragraphs[0]
        kp.paragraph_format.space_after = Pt(0)
        _run(kp, kicker, bold=True, size=7.5, color=kicker_color, caps=True)
    tp = cell.add_paragraph()
    tp.paragraph_format.space_before = Pt(1); tp.paragraph_format.space_after = Pt(0)
    _run(tp, title, bold=True, size=title_size, color=title_color, font=SERIF)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return t


def _header(doc, right_title, grade, subject, dt, title_color=PINE):
    t = doc.add_table(rows=1, cols=2)
    _no_borders(t)
    lc = t.cell(0, 0).paragraphs[0]
    r = _run(lc, "Aruvi", bold=True, size=17, color=PINE, font=SERIF)
    _run(lc, ".", bold=True, size=17, color=CLAY, font=SERIF)
    _run(lc, "  LESSON STUDIO", size=7.5, color=SOFT, caps=True)
    lc2 = t.cell(0, 0).add_paragraph()
    _run(lc2, "NCF 2023 aligned", italic=True, size=8, color=SOFT, font=SERIF)
    rc = t.cell(0, 1).paragraphs[0]
    rc.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    _run(rc, right_title, bold=True, size=12, color=title_color, font=SERIF)
    rc2 = t.cell(0, 1).add_paragraph()
    rc2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    _run(rc2, f"Grade {grade} · {subject} · {_date(dt)}", size=8, color=SOFT)
    # heavy rule under the header
    rule = doc.add_table(rows=1, cols=1)
    b = OxmlElement("w:tblBorders")
    e = OxmlElement("w:bottom")
    e.set(qn("w:val"), "single"); e.set(qn("w:sz"), "18"); e.set(qn("w:space"), "0"); e.set(qn("w:color"), "1A1917")
    b.append(e)
    rule._tbl.tblPr.append(b)
    rule.cell(0, 0).paragraphs[0].paragraph_format.space_after = Pt(0)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)


def _date(dt: datetime) -> str:
    return f"{dt.day} {dt.strftime('%B')} {dt.year}"


def _grid_table(doc, headers, values, widths=None):
    """A bordered header+value grid (metadata / competency / stage tables)."""
    t = doc.add_table(rows=2, cols=len(headers))
    _hairlines(t)
    for i, h in enumerate(headers):
        c = t.cell(0, i); _bg(c, "F1ECE2")
        p = c.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _run(p, h, bold=True, size=7.5, color=SOFT, caps=True)
    for i, v in enumerate(values):
        c = t.cell(1, i)
        p = c.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _run(p, v, size=9)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return t


def _setup(doc):
    from docx.shared import Inches
    for s in doc.sections:
        s.top_margin = Inches(0.6); s.bottom_margin = Inches(0.6)
        s.left_margin = Inches(0.7); s.right_margin = Inches(0.7)
    st = doc.styles["Normal"]
    st.font.name = SANS; st.font.size = Pt(9.5)


# ── lesson-plan blocks ───────────────────────────────────────────────────────

def _lp_front(doc, lp, competencies, spines=None):
    ch_no = lp.get("chapter_number", ""); ch_title = lp.get("chapter_title", "")
    groups = lp.get("groups", []) or []
    total_periods = lp.get("total_periods") or sum(len(g.get("periods", []) or []) for g in groups)
    _grid_table(doc, ["Chapter", "Title", "Total periods", "Duration"],
                [f"Ch {ch_no}", ch_title, str(total_periods), _duration_breakdown(lp)])
    if spines:
        # English's 4-column spine table (Spine · Section Name · Code · Competency), spine +
        # section cells vertically merged across their competency rows.
        total = sum(len(s.get("rows", [])) for s in spines)
        t = doc.add_table(rows=1 + total, cols=4); _hairlines(t)
        for j, h in enumerate(["Spine", "Section Name", "Code", "Competency"]):
            c = t.cell(0, j); _bg(c, "F1ECE2"); _run(c.paragraphs[0], h, bold=True, size=7.5, color=SOFT, caps=True)
        ri = 1
        for s in spines:
            rows = s.get("rows", [])
            start = ri
            for k, row in enumerate(rows):
                if k == 0:
                    _run(t.cell(ri, 0).paragraphs[0], s.get("spine", ""), bold=True, size=9, color=INK)
                    _run(t.cell(ri, 1).paragraphs[0], s.get("section", ""), size=9, color=BODY)
                _run(t.cell(ri, 2).paragraphs[0], row.get("c_code", ""), bold=True, size=9, color=PINE)
                _run(t.cell(ri, 3).paragraphs[0], row.get("text", ""), size=9, color=BODY)
                ri += 1
            if len(rows) > 1:
                t.cell(start, 0).merge(t.cell(start + len(rows) - 1, 0))
                t.cell(start, 1).merge(t.cell(start + len(rows) - 1, 1))
        doc.add_paragraph().paragraph_format.space_after = Pt(2)
    elif competencies:
        has_just = any(c.get("justification") for c in competencies)
        cols = ["C No.", "Targeted NCF competencies"] + (["Justification"] if has_just else [])
        t = doc.add_table(rows=1 + len(competencies), cols=len(cols)); _hairlines(t)
        for j, h in enumerate(cols):
            c = t.cell(0, j); _bg(c, "F1ECE2"); _run(c.paragraphs[0], h, bold=True, size=7.5, color=SOFT, caps=True)
        for i, comp in enumerate(competencies, 1):
            cc = t.cell(i, 0); pp = cc.paragraphs[0]; pp.alignment = WD_ALIGN_PARAGRAPH.CENTER
            _run(pp, comp.get("c_code", ""), bold=True, size=9, color=PINE)
            _run(t.cell(i, 1).paragraphs[0], comp.get("text", ""), size=9, color=BODY)
            if has_just:  # chapter-level rationale in a third column, same font/size as the competency
                _run(t.cell(i, 2).paragraphs[0], comp.get("justification", ""), size=9, color=BODY)
        doc.add_paragraph().paragraph_format.space_after = Pt(2)
    # English's spine table already lists the spines, so skip the redundant progression table there.
    # Mathematics has no meaningful progression axis, so skip it for maths too (founder 2026-07-20).
    # Secondary Science is section-anchored and FLAT (no per-period progression_stage — see science
    # subject.py), so its progression table would only re-list the chapter sections; skip it there
    # too (founder 2026-07-20). Detect by the data, not the grade string: secondary science emits
    # top-level groups typed "section", middle science "progression_stage"/"stage". Social Sciences'
    # v3 edge-model plan collapses to ONE flat "unit" group ("Units"), so its progression table
    # would print a single meaningless "Unit 1 · Units" row — skip it there too (founder 2026-07-20).
    # Shared by both the plain LP and integrated DOCX (mirrors export_lesson_pdf._no_progression).
    _science_sectioned = (
        lp.get("subject", "") == "science" and bool(groups) and groups[0].get("type") == "section"
    )
    _ss_flat_units = (
        lp.get("subject", "") == "social_sciences" and bool(groups) and groups[0].get("type") == "unit"
    )
    # The World Around Us (prep only) is section-anchored with no progression axis, so its
    # progression table would just re-list the chapter sections — skip it there too (founder 2026-07-20).
    _twau = lp.get("subject", "") == "the_world_around_us"
    if (groups and not spines and lp.get("subject", "") != "mathematics"
            and not _science_sectioned and not _ss_flat_units and not _twau):
        word = _lp_group_word(groups)
        t = doc.add_table(rows=1 + len(groups) + 1, cols=2); _hairlines(t)
        _bg(t.cell(0, 0), "F1ECE2"); _run(t.cell(0, 0).paragraphs[0], f"{word} No.", bold=True, size=7.5, color=SOFT, caps=True)
        _bg(t.cell(0, 1), "F1ECE2"); _run(t.cell(0, 1).paragraphs[0], f"Progression {word}", bold=True, size=7.5, color=SOFT, caps=True)
        for i, g in enumerate(groups, 1):
            pp = t.cell(i, 0).paragraphs[0]; pp.alignment = WD_ALIGN_PARAGRAPH.CENTER
            _run(pp, str(i), bold=True, size=9, color=PINE)
            _run(t.cell(i, 1).paragraphs[0], g.get("label", ""), size=9, color=BODY)
        blurb = AXIS_INFO.get(groups[0].get("type", ""))
        note = t.cell(len(groups) + 1, 0)
        note.merge(t.cell(len(groups) + 1, 1))
        if blurb:
            _run(note.paragraphs[0], blurb, italic=True, size=8, color=SOFT)
        doc.add_paragraph().paragraph_format.space_after = Pt(2)


def _lp_unit(doc, p, *, first_pedagogy):
    meta = p.get("meta", {}) or {}
    dur = meta.get("duration_minutes")
    head = _para(doc, space_before=8, space_after=2)
    _run(head, f"Period {p.get('number')}  ", bold=True, size=10.5, color=INK)
    if dur:
        _run(head, f"{dur} min   ", size=9, color=SOFT)
    _run(head, p.get("title", ""), bold=True, size=10.5, color=INK, font=SERIF)
    ped = p.get("approach") or ""
    if ped:
        pp = _para(doc, space_after=2)
        if first_pedagogy:
            _run(pp, "Pedagogy: ", bold=True, size=10.5, color=INK)
        _run(pp, ped, italic=True, size=10.5, color=SOFT)
    mats = p.get("materials") or []
    if mats:
        mp = _para(doc, space_after=2)
        _run(mp, "Materials: ", bold=True, size=9, color=INK)
        _run(mp, ", ".join(mats), size=9, color=BODY)
    notes = p.get("teacher_notes") or []
    if notes:
        npp = _para(doc, space_after=3)
        _run(npp, "Teacher notes: ", bold=True, size=9, color=INK)
        _run(npp, " ".join(notes), size=9, color=BODY)
    phases = [ph for ph in (p.get("phases") or []) if ph.get("text") or ph.get("label")]
    if phases:
        t = doc.add_table(rows=len(phases), cols=2); _hairlines(t)
        # Fixed widths so the time column is just wide enough and every period's table lines up:
        # a narrow time column, the rest of the page for the phase text.
        sec = doc.sections[0]
        avail = (sec.page_width - sec.left_margin - sec.right_margin) / 914400  # EMU → inches
        time_in = 0.6
        _fixed_table(t, [time_in, round(avail - time_in, 2)])
        for i, ph in enumerate(phases):
            _run(t.cell(i, 0).paragraphs[0], _phase_duration(ph), size=8, color=SOFT)
            _run(t.cell(i, 1).paragraphs[0], ph.get("text", ""), size=9, color=BODY)
    elif p.get("activities"):
        for a in p["activities"]:
            _run(_para(doc, space_after=1), a, size=9, color=BODY)
    hw = p.get("homework")
    if hw:
        hp = _para(doc, space_before=2, space_after=2)
        _run(hp, "Homework: ", bold=True, size=9, color=INK)
        _run(hp, str(hw).replace("**", ""), size=9, color=BODY)


# ── assessment blocks ────────────────────────────────────────────────────────

def _stimulus(doc, block):
    if not block:
        return
    if block.get("type") == "table" and block.get("table"):
        tb = block["table"]; header = tb.get("header", []) or []; rows = tb.get("rows", []) or []
        ncols = max([len(header)] + [len(r) for r in rows]) if (header or rows) else 0
        if ncols:
            t = doc.add_table(rows=(1 if header else 0) + len(rows), cols=ncols); _hairlines(t, "CDE0D8")
            ri = 0
            if header:  # a real data table — the first row is a filled/bold header
                for j, h in enumerate(header):
                    c = t.cell(0, j); _bg(c, GHEAD_HEX); _run(c.paragraphs[0], h, bold=True, size=8, color=G_DARK)
                ri = 1
            for row in rows:  # word bank (no header) → all cells same plain style
                for j, val in enumerate(row):
                    if j < ncols:
                        _run(t.cell(ri, j).paragraphs[0], val, size=8, color=BODY)
                ri += 1
            return
    content = block.get("content")
    if content and block.get("type") != "svg":
        _run(_para(doc, space_after=2), content, italic=True, size=9, color=BODY)
    elif block.get("type") == "svg":
        _run(_para(doc, space_after=2), "[figure — see the on-screen version]", italic=True, size=8, color=SOFT)


def _assess_item(doc, it, qn, include_answers):
    n = it.get("normalized") or {}
    lo = n.get("linked_lo") or it.get("implied_lo") or ""
    cog = n.get("cognitive_demand")
    stem = n.get("stem") or it.get("prompt") or ""
    template = n.get("template") or ""
    tf = n.get("tf_statements") or []
    is_tf = template == "true_false" and bool(tf)
    stem_parts = n.get("stem_parts") or []
    stem_lead = n.get("stem_lead") or ""
    # For TRUE_FALSE / a parsed multi-part stem, the header line is just the instruction (stem_lead);
    # the statements/parts print ONCE below (never also as options) — matches the on-screen view.
    body_stem = stem_lead if ((is_tf or stem_parts) and stem_lead) else stem
    # green question header block
    if lo or cog:
        t = doc.add_table(rows=1, cols=2); _no_borders(t)
        # Fixed widths so Word doesn't starve the learning-outcome column: LO keeps ~3/4 of
        # the page, the cognitive-demand column a narrow strip at the far right (matches the PDF).
        sec = doc.sections[0]
        avail = (sec.page_width - sec.left_margin - sec.right_margin) / 914400  # EMU → inches
        cog_in = round(avail * 0.24, 2)
        _fixed_table(t, [round(avail - cog_in, 2), cog_in])
        lcell = t.cell(0, 0); _bg(lcell, GHEAD_HEX)
        lp0 = lcell.paragraphs[0]; lp0.paragraph_format.space_after = Pt(0)
        _run(lp0, f"Q{qn}", bold=True, size=10, color=G_DARK, font=SERIF)
        if lo:
            _run(lp0, "  ·  LEARNING OUTCOME", bold=True, size=7, color=G_ACCENT)
            lp1 = lcell.add_paragraph(); lp1.paragraph_format.space_before = Pt(0)
            _run(lp1, lo, size=9, color=RGBColor(0x1F, 0x3A, 0x30))
        rcell = t.cell(0, 1); _bg(rcell, GHEAD_HEX)
        rp0 = rcell.paragraphs[0]; rp0.alignment = WD_ALIGN_PARAGRAPH.RIGHT; rp0.paragraph_format.space_after = Pt(0)
        if cog:
            _run(rp0, "COGNITIVE DEMAND", bold=True, size=7, color=G_ACCENT)
            rp1 = rcell.add_paragraph(); rp1.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            _run(rp1, cog, size=9, color=RGBColor(0x1F, 0x3A, 0x30))
        sp = _para(doc, space_before=3, space_after=2)
        _run(sp, body_stem, size=9.5, color=INK)
    else:
        sp = _para(doc, space_before=6, space_after=2)
        _run(sp, f"Q{qn}. ", bold=True, size=9.5, color=INK)
        _run(sp, body_stem, size=9.5, color=INK)
    # stimulus
    _stimulus(doc, n.get("passage"))
    if template != "passage":
        _stimulus(doc, n.get("visual_stimulus"))
    if it.get("visual_stimulus") and not (n.get("visual_stimulus") or n.get("passage")):
        vs = it["visual_stimulus"]
        if isinstance(vs, dict):
            _stimulus(doc, vs)
    # TRUE_FALSE statements / multi-part stems — printed once
    if is_tf:
        for s in tf:
            op = _para(doc, space_after=1)
            _run(op, f"{s.get('marker')}  ", bold=True, size=9, color=BODY)
            _run(op, s.get("text", ""), size=9, color=BODY)
    elif stem_parts:
        for p in stem_parts:
            op = _para(doc, space_after=1)
            _run(op, f"{p.get('marker')}  ", bold=True, size=9, color=BODY)
            _run(op, p.get("text", ""), size=9, color=BODY)
    # options (never for TRUE_FALSE)
    if not is_tf:
        for o in (n.get("options") or []):
            correct = include_answers and o.get("is_correct")
            op = _para(doc, space_after=1)
            _run(op, f"{o.get('label')}.  ", bold=True, size=9, color=(G_ACCENT if correct else BODY))
            _run(op, o.get("text", "") + (" ✓" if correct else ""), size=9,
                 color=(G_ACCENT if correct else BODY), bold=bool(correct))
    for ln in (n.get("scaffold_lines") or []):
        if ln:
            _run(_para(doc, space_after=1), ln, size=9, color=BODY)
    if include_answers:
        _answer_block(doc, n)


def _answer_block(doc, n):
    # TRUE_FALSE: per-statement verdict (+ reason) — replaces the correct-answer list.
    tf = n.get("tf_statements") or []
    if (n.get("template") == "true_false") and tf:
        hp = _para(doc, space_before=4, space_after=2, align=WD_ALIGN_PARAGRAPH.CENTER)
        _bg_para(hp, GTINT_HEX); _run(hp, "Answer section", bold=True, size=10, color=G_DARK)
        if not any(s.get("reason") for s in tf) and n.get("model_answer"):
            p = _para(doc, space_after=1)
            _run(p, "Suggested answer: ", bold=True, size=8.5, color=G_DARK)
            _run(p, n.get("model_answer"), size=9, color=RGBColor(0x1F, 0x3A, 0x30))
        else:
            _run(_para(doc, space_after=1), "Answer key", bold=True, size=8.5, color=G_DARK)
            for s in tf:
                p = _para(doc, space_after=1)
                _run(p, f"{s.get('marker')}  ", bold=True, size=9, color=G_ACCENT)
                _run(p, "True" if s.get("verdict") else "False", bold=True, size=9, color=RGBColor(0x1F, 0x3A, 0x30))
                if s.get("reason"):
                    _run(p, f" — {s.get('reason')}", size=9, color=RGBColor(0x1F, 0x3A, 0x30))
        return
    bits = []
    correct = [o.get("label") for o in (n.get("options") or []) if o.get("is_correct")]
    if correct:
        bits.append(("Correct answer:", ", ".join(str(c) for c in correct)))
    if n.get("model_answer"):
        bits.append(("Model answer:", n.get("model_answer")))
    reveals = n.get("option_reveals") or {}
    method = n.get("method_one_line")
    elems = n.get("expected_elements") or []
    looks = n.get("look_fors") or []
    ex = n.get("exercise_ref") or n.get("exercise_desc")
    if not (bits or reveals or method or elems or looks or ex):
        return
    hp = _para(doc, space_before=4, space_after=2, align=WD_ALIGN_PARAGRAPH.CENTER)
    _bg_para(hp, GTINT_HEX)
    _run(hp, "Answer section", bold=True, size=10, color=G_DARK)
    for k, v in bits:
        p = _para(doc, space_after=1)
        _run(p, k + " ", bold=True, size=8.5, color=G_DARK); _run(p, v, size=9, color=RGBColor(0x1F, 0x3A, 0x30))
    if reveals:
        _run(_para(doc, space_after=1), "What each choice reveals", bold=True, size=8.5, color=G_DARK)
        for lab, txt in reveals.items():
            p = _para(doc, space_after=1)
            if lab != "note":
                _run(p, f"{lab}  ", bold=True, size=8.5, color=G_ACCENT)
            _run(p, txt, size=9, color=RGBColor(0x1F, 0x3A, 0x30))
    for label, items in (("Expected elements", elems), ("Look for", looks)):
        if items:
            _run(_para(doc, space_after=1), label, bold=True, size=8.5, color=G_DARK)
            for x in items:
                _run(_para(doc, space_after=0), "• " + x, size=9, color=RGBColor(0x1F, 0x3A, 0x30))
    if method:
        p = _para(doc, space_after=1); _run(p, "Method: ", bold=True, size=8.5, color=G_DARK); _run(p, method, size=9, color=RGBColor(0x1F, 0x3A, 0x30))
    if ex:
        p = _para(doc, space_after=1); _run(p, "Textbook exercise: ", bold=True, size=8.5, color=G_DARK)
        _run(p, f"{n.get('exercise_ref') or ''} — {n.get('exercise_desc') or ''}".strip(" —"), size=9, color=RGBColor(0x1F, 0x3A, 0x30))


def _bg_para(p, hex_color):
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear"); shd.set(qn("w:fill"), hex_color)
    p._p.get_or_add_pPr().append(shd)


# ── public entry points ──────────────────────────────────────────────────────

def export_lesson_plan_docx(view, *, competencies=None, competency_spines=None, plan_date=None, generated_at=None) -> bytes:
    lp = view.get("lesson_plan", {}) or {}
    doc = Document(); _setup(doc)
    _header(doc, "Lesson plan", grade_roman(lp.get("grade", "")), subject_display(lp.get("subject", "")),
            generated_at or plan_date or datetime.now())
    _lp_front(doc, lp, competencies or [], spines=competency_spines)
    _groups = lp.get("groups", []) or []
    word = _lp_group_word(_groups)
    # Mathematics renders its units flat — no section/stage band (founder 2026-07-20). Social
    # Sciences' v3 edge-model plan is likewise ONE flat "unit" group ("Units") with no axis — its
    # band would read a bare "UNIT 1 · Units" over every period — so suppress it there too.
    _ss_flat_units = (
        lp.get("subject", "") == "social_sciences" and bool(_groups) and _groups[0].get("type") == "unit"
    )
    _show_band = lp.get("subject", "") != "mathematics" and not _ss_flat_units
    first = True
    for i, g in enumerate(lp.get("groups", []) or [], 1):
        if _show_band:
            _shaded_band(doc, KRAFT_HEX, f"{word.upper()} {i}", PINE, g.get("label", ""), INK)
        for p in _iter_group_periods(g):
            _lp_unit(doc, p, first_pedagogy=first); first = False
    return _tobytes(doc)


def export_assessment_docx(view, *, include_answers=False, assessment_type="Formative", plan_date=None, generated_at=None) -> bytes:
    av = view.get("assessment", {}) or {}
    groups = av.get("groups", []) or []
    doc = Document(); _setup(doc)
    _header(doc, "Chapter Assessment", grade_roman(av.get("grade", "")), subject_display(av.get("subject", "")),
            generated_at or plan_date or datetime.now(), title_color=G_ACCENT)
    total_q = sum(len(g.get("items", []) or []) for g in groups)
    _grid_table(doc, ["Chapter", "Title", "Total questions", "Type"],
                [f"Ch {av.get('chapter_number','')}", av.get("chapter_title", ""), str(total_q), assessment_type])
    _run(_para(doc, space_after=4), intro_paragraph(av.get("subject", ""), groups, av.get("grade")), size=9, color=BODY)
    qn = 0
    for i, g in enumerate(groups, 1):
        gtype = g.get("type", "")
        if gtype == "question_type":
            # SS / TWAU group by question type — "QUESTION TYPE" kicker (no number) over the FULL
            # type name, not the raw "STAGE n / MCQ" the generic path produced (founder 2026-07-20).
            kicker = "QUESTION TYPE"
            title = question_type_full(g.get("label"))
        else:
            word = CTX_WORD.get(gtype, "Stage")
            kicker = f"{word.upper()} {(g.get('meta',{}) or {}).get('stage_number', i)}"
            title = g.get("label", "")
        _shaded_band(doc, GHEAD_HEX, kicker, G_ACCENT, title, G_DARK, title_size=11)
        for it in g.get("items", []) or []:
            qn += 1
            _assess_item(doc, it, qn, include_answers)
    return _tobytes(doc)


def export_integrated_docx(view, *, include_answers=False, unit_number=None, competencies=None, competency_spines=None, plan_date=None, generated_at=None) -> bytes:
    lp = view.get("lesson_plan", {}) or {}
    av = view.get("assessment", {}) or {}
    groups = lp.get("groups", []) or []
    doc = Document(); _setup(doc)
    _header(doc, "Lesson plan & assessment", grade_roman(lp.get("grade", "")), subject_display(lp.get("subject", "")),
            generated_at or plan_date or datetime.now())
    if unit_number is None:
        _lp_front(doc, lp, competencies or [], spines=competency_spines)
    _run(_para(doc, space_after=4), intro_paragraph(lp.get("subject", ""), groups, lp.get("grade")), size=9, color=BODY)
    by_anchor = _items_by_anchor(av)
    word = _lp_group_word(groups)
    # Mathematics renders its units flat — no section/stage band (founder 2026-07-20). Social
    # Sciences' v3 edge-model plan is likewise ONE flat "unit" group ("Units") with no axis — its
    # band would read a bare "UNIT 1 · Units" over every period — so suppress it there too.
    _ss_flat_units = (
        lp.get("subject", "") == "social_sciences" and bool(groups) and groups[0].get("type") == "unit"
    )
    _show_band = lp.get("subject", "") != "mathematics" and not _ss_flat_units
    qn = 0; first = True
    for i, g in enumerate(groups, 1):
        periods = _iter_group_periods(g)
        if unit_number is not None and not any(p.get("number") == unit_number for p in periods):
            continue
        if _show_band:
            _shaded_band(doc, KRAFT_HEX, f"{word.upper()} {i}", PINE, g.get("label", ""), INK)
        for p in periods:
            if unit_number is not None and p.get("number") != unit_number:
                continue
            _lp_unit(doc, p, first_pedagogy=first); first = False
            items = by_anchor.get(p.get("number"), [])
            if items:
                ap = _para(doc, space_before=8, space_after=3, align=WD_ALIGN_PARAGRAPH.CENTER)
                _run(ap, "Assessment", bold=True, size=11, color=G_DARK)
            for it in items:
                qn += 1
                _assess_item(doc, it, qn, include_answers)
    return _tobytes(doc)


def _tobytes(doc) -> bytes:
    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()
