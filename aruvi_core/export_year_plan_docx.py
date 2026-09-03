"""
Year Plan — DOCX export of the Year Plan TABLE (and only the table).

Why this module exists at all, given YearPlan.jsx's own header says the year plan is
"a LIVING mobile view, never a PDF": that decision was about the ALLOCATION REPORT —
a document the teacher was asked to read INSTEAD of the screen. This is the opposite
direction. She reads the year on the screen and then needs the same three columns
somewhere the screen cannot go: a staff meeting, a HOD's file, a printout on a
noticeboard. So the export is deliberately the table, the header that identifies it
and the note that explains the two columns — never a report built around it. No
competencies, no effort-index values, no summary strip, no executive prose. If it is
not on the pane, it is not in the file.

★ THE PAYLOAD IS THE SCREEN'S OWN MODEL, NOT A RE-COMPUTATION.

The suggested-periods column is computed CLIENT-side (YearPlan.jsx's `useMemo`:
budget from the readiness projection, distributed by chapter weight with
largest-remainder). Recomputing it here would put a second implementation of that
arithmetic in the product, and the day the two drift the teacher has a Word document
that contradicts the screen she exported it from — the exact defect the calibrated-
budget work of 2026-08-21 was about (Year Plan said 14 where the chapter step said
19). So the client POSTs what it is displaying and this module only renders it. That
is also why every value arrives pre-formatted-or-null: a missing number renders as an
em-dash, NEVER as a zero (the Support `metaErr` rule — a document may say it does not
know; it may never invent an answer about her record).

House style is IMPORTED from export_allocation_docx, not copied, so the three Aruvi
documents cannot drift (the convention export_data_rights_docx.py set).

Payload shape (see api/main.YearPlanExportRequest):
    {subject: "Science" (display) | "science" (slug), grade: "vii",
     budget: int|None, generated_at: iso|None,
     rows: [{n: int, title: str, sug: int|None, plan: int|None,
             prepared: bool, awaited: bool}],
     sug_total: int, plan_total: int}
"""

from __future__ import annotations

from datetime import datetime
from io import BytesIO
from typing import Any, Dict, List, Optional

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from . import brand   # the MEYY wordmark raster (2026-09-03)

from .export_allocation_docx import (
    _bg, _cell, _hairlines, _no_borders, _rule, _run, _set_widths,
    INK, PINE, CLAY, GRAY, MUTED, BODY, SERIF, ALT_HEX, TOTAL_HEX,
)
from .report_competency import grade_roman, subject_display, date_long

# The em-dash the screen uses for "no number here" (.yp-dash). One constant so the
# three places that need it cannot disagree.
DASH = "—"


def _fmt(v: Optional[int]) -> str:
    """A number, or the em-dash. Zero is a real answer and prints as 0; None is not."""
    return DASH if v is None else str(v)


def _plan_cell(row: Dict[str, Any]) -> str:
    """The "Your plan" cell, mirroring YearPlan.jsx's four-way render exactly.

    A chapter can be prepared with no periods recorded (legacy prepares stored no
    `prepared_periods`) — the screen shows "set" for those, and so does this, because
    a dash there would read as "not prepared" and misstate her year.
    """
    if row.get("awaited"):
        return DASH
    if row.get("plan") is not None:
        return str(row["plan"])
    if row.get("prepared"):
        return "set"
    return DASH


def _when(generated_at: Optional[str]) -> datetime:
    if generated_at:
        try:
            return datetime.fromisoformat(generated_at)
        except ValueError:
            pass
    return datetime.now()


def export_year_plan_docx(payload: Dict[str, Any]) -> bytes:
    """Render the Year Plan table as an editable Word document. Returns docx bytes."""
    rows: List[Dict[str, Any]] = list(payload.get("rows") or [])
    subject = payload.get("subject") or ""
    grade = payload.get("grade") or ""
    budget = payload.get("budget")

    # Display strings. `subject` may arrive as a slug ("social_sciences") or as the
    # profile's display name ("Social Sciences"); subject_display normalizes either.
    subj = subject_display(subject)
    g = grade_roman(grade)

    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = Inches(0.7)
    sec.bottom_margin = Inches(0.7)
    sec.left_margin = Inches(0.7)
    sec.right_margin = Inches(0.7)

    # ── Header — the same mark, rule and right-hand identification block as the
    # allocation report and the data export. Only the title differs.
    ht = doc.add_table(rows=1, cols=2)
    _no_borders(ht)
    ht.columns[0].width = Inches(4.0)
    ht.columns[1].width = Inches(3.1)
    lc = ht.rows[0].cells[0]
    lc.paragraphs[0].clear()
    lp = lc.paragraphs[0]
    lp.paragraph_format.space_after = Pt(0)
    brand.add_wordmark(lp, 17)   # the MEYY wordmark (2026-09-03)
    lp2 = lc.add_paragraph()   # kicker beneath the mark; NCF line gone (founder, 2026-09-03)
    lp2.paragraph_format.space_before = Pt(1)
    _run(lp2, "LESSON STUDIO", size=7, color=GRAY, font="Calibri")

    rc = ht.rows[0].cells[1]
    rc.paragraphs[0].clear()
    rp = rc.paragraphs[0]
    rp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    rp.paragraph_format.space_after = Pt(0)
    _run(rp, "Year plan", size=12, bold=True, color=PINE, font=SERIF)
    rp2 = rc.add_paragraph()
    rp2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    # "Class VII" — the teacher-facing word on this pane is Class, not Grade (the web
    # side says "Class 7"); the allocation report says Grade because that document
    # predates the change and its own wording is settled.
    _run(rp2, f"Class {g} · {subj} · {date_long(_when(payload.get('generated_at')))}",
         size=8, color=GRAY)

    gap = doc.add_paragraph()
    gap.paragraph_format.space_after = Pt(8)
    _rule(gap, color="1A1917", sz=16)

    _year_plan_table(doc, rows, payload)
    _note(doc, rows, budget)

    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.getvalue()


def _year_plan_table(doc, rows: List[Dict[str, Any]], payload: Dict[str, Any]) -> None:
    """The three columns of the pane: Chapter · Suggested periods · Your plan.

    Column ORDER and wording are the screen's, deliberately — a teacher who exports
    what she is looking at should be able to lay the two side by side.
    """
    headers = ["#", "Chapter", "Suggested periods", "Your plan"]
    t = doc.add_table(rows=1, cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    _hairlines(t, "DDDDDD")
    for i, h in enumerate(headers):
        _cell(t.rows[0].cells[i], h.upper(), size=6.5, color=GRAY,
              align="left" if i == 1 else "center")

    for idx, r in enumerate(rows, 1):
        cells = t.add_row().cells
        _cell(cells[0], str(r.get("n", idx)).zfill(2), size=7, color=MUTED, align="center")
        title = r.get("title") or ""
        if r.get("awaited"):
            # "Book awaited" chapters hold periods in the budget but can carry no plan
            # (API flag, 2026-08-06). They belong in the year; the screen greys them
            # and so does this — MUTED ink rather than a note that would need a legend.
            _cell(cells[1], title, size=8, color=MUTED, align="left")
        else:
            _cell(cells[1], title, size=8, align="left")
        _cell(cells[2], _fmt(r.get("sug")), size=8, align="center")
        _cell(cells[3], _plan_cell(r), size=8, align="center",
              color=MUTED if _plan_cell(r) == DASH else INK)
        if idx % 2 == 0:
            for cell in cells:
                _bg(cell, ALT_HEX)

    # Total row — the analogue of .yp-tot, and it carries the SCREEN's totals rather
    # than a sum of the column above it. They agree today; if a future row type ever
    # counts differently, the document must still say what the pane said.
    cells = t.add_row().cells
    _cell(cells[0], "", size=8)
    _cell(cells[1], "Total periods", size=8, bold=True, align="left")
    _cell(cells[2], _fmt(payload.get("sug_total")), size=8, bold=True, align="center")
    _cell(cells[3], _fmt(payload.get("plan_total")), size=8, bold=True, align="center")
    for cell in cells:
        _bg(cell, TOTAL_HEX)

    # ≈7.1in of content width at 0.7in margins.
    _set_widths(t, [0.4, 3.9, 1.5, 1.3])


def _note(doc, rows: List[Dict[str, Any]], budget: Optional[int]) -> None:
    """The pane's own explanatory note — ALL of it, and in the pane's own words.

    It is here because the two period columns are not self-explanatory off the screen:
    a colleague reading the printout has no way to know that "Suggested" is Aruvi's
    proposal and "Your plan" is the teacher's own commitment.

    ⚠️ THIS PROSE IS A SECOND COPY OF `.yp-note` IN YearPlan.jsx — change one, change
    the other. The closing "To know how Aruvi suggests…" sentence was dropped on the
    first build (founder, live) precisely because nothing tied the two together, so
    `test_note_matches_the_panes_own_words` now reads the JSX and checks every
    non-interpolated sentence of the pane's note appears here. Splitting is not worth
    a shared string table across the language boundary; the test is.
    """
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(4)
    _run(p, "Your teaching year at a glance — how ", size=9, color=BODY, font=SERIF)
    if budget is not None:
        _run(p, "a budget of ", size=9, color=BODY, font=SERIF)
        _run(p, f"{budget} periods", size=9, bold=True, color=INK, font=SERIF)
    else:
        _run(p, "your periods", size=9, color=BODY, font=SERIF)
    _run(p, f" spread across all {len(rows)} chapters. ", size=9, color=BODY, font=SERIF)
    _run(p, "Suggested periods", size=9, bold=True, color=INK, font=SERIF)
    _run(p, " is Meyy's proposal, giving heavier chapters more room. Each time you "
            "prepare a lesson you set your own periods for that chapter; those appear in ",
         size=9, color=BODY, font=SERIF)
    _run(p, "Your plan", size=9, bold=True, color=INK, font=SERIF)
    _run(p, ", beside the suggestion, so you can see where you've adjusted and how much "
            "of the year you've committed. To know how Meyy suggests, refer to Ask Meyy "
            "time allocation section.", size=9, color=BODY, font=SERIF)
