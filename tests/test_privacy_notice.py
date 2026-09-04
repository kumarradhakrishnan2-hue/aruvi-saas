"""
The Privacy Notice — served, shown, recorded; never signed (2026-09-04).

Pinned here on purpose:
  * the notice is a document family of its own (privacy_policy_v{V}.md) and the consent
    parser does not see it — adding the notice beside the agreement must not change
    which agreement version is current;
  * the notice has no ticks — it is GIVEN (DPDP §5), and a consent screen grown out of
    it would be the wrong instrument;
  * GET /legal/privacy needs NO identity — the sign-in screen links it before any
    account exists;
  * the version SHOWN is stamped on the account at registration and on dismissal of the
    "updated" note, rendered in her export, and gone with her account;
  * ★ the erasure receipt's `_KEPT` and the notice's §7 name the SAME things — the
    receipt is the one document she holds after her account is gone, so it may not name
    less than the notice promised, nor more. Sabotage-verified: remove a row from either
    side and this fails.

Run standalone:  python3 tests/test_privacy_notice.py     (also pytest-compatible)
"""
from __future__ import annotations

import io
import os
import re
import sys
import tempfile

sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.abspath(__file__))))

_TMP_STATE = tempfile.mkdtemp(prefix="aruvi-test-state-")
os.environ.setdefault("ARUVI_STATE_DIR", _TMP_STATE)

from api import legal  # noqa: E402
from aruvi_core.adapters.data_rights_service_file import _KEPT  # noqa: E402


def _section_7(body: str) -> str:
    """The text of '## 7. How long we keep it…' up to the next '## '."""
    m = re.search(r"^## 7\..*?(?=^## 8\.)", body, re.S | re.M)
    assert m, "the notice must keep a §7 retention section and a §8 after it"
    return m.group(0)


def test_notice_loads_and_is_not_a_consent_document():
    doc = legal.load_privacy_document()
    assert doc["document_id"] == "privacy_policy"
    # "Privacy Notice", bare (founder, 2026-09-04: not "Meyy — Privacy Notice" — the bar
    # above it already says whose it is).
    assert doc["title"] == "Privacy Notice"
    assert "(Draft" not in doc["title"], "the draft marker belongs to the file, not the screen"
    assert doc["version"] == legal.current_privacy_version()
    assert re.match(r"\d{4}-\d{2}-\d{2}$", doc["published"]), "dated footer parsed"
    body = doc["body"]
    assert "☐" not in body, "a notice has no ticks — it is given, not signed"
    assert not any(ln.lstrip().startswith(">") for ln in body.splitlines()), \
        "front matter (the note to the lawyer) must never reach a teacher"
    assert "For legal review" not in body, "the dated footer line is lifted out of the body"
    assert "## 2." in body and "## 8." in body and "## 10." in body
    print("✓ Notice loads: title, version, date, body — and no ticks")


def test_consent_parser_ignores_the_notice():
    """Dropping privacy_policy_v0.1.md beside consent_and_disclaimer_v0.4.md must not
    make '0.1' the current AGREEMENT — or every signer would be re-ticked against a
    document that has no acknowledgements and the route would 503."""
    assert "privacy_policy" not in " ".join(legal.available_versions())
    cur = legal.current_version()
    assert legal.load_consent_document(cur)["document_id"] == "consent_and_disclaimer"
    assert len(legal.load_consent_document(cur)["acknowledgements"]) == 5
    print(f"✓ Consent parser unaffected — agreement v{cur} still current, five ticks")


# What each receipt row promises, as a phrase that must appear in the notice's §7.
_ROW_PHRASE = {
    "Disaster-recovery backups": "Disaster-recovery backups",
    "Tax records for payments made": "Your invoices",
    "Email we exchanged": "Email we exchanged",
    "The record that you asked us to erase": "The record that you asked us to erase",
    "Your acceptance of the User Agreement": "The record that you accepted the User Agreement",
    "Shared lesson-plan library content": "Shared lesson-plan library",
}


def test_receipt_kept_rows_match_notice_section_7():
    """★ Four places say what survives an erasure and must agree; this pins two of them
    to each other — the receipt (`_KEPT`) and the notice (§7). Both directions."""
    s7 = _section_7(legal.load_privacy_document()["body"])
    whats = [k["what"] for k in _KEPT]
    # Every receipt row is promised in the notice…
    for what in whats:
        assert what in _ROW_PHRASE, (
            f"_KEPT gained a row ({what!r}) the notice does not know — add it to §7 of "
            "the notice AND to _ROW_PHRASE here, or the receipt names more than she was told")
        assert _ROW_PHRASE[what] in s7, (
            f"the notice's §7 no longer states {what!r} ({_ROW_PHRASE[what]!r}) — "
            "the receipt promises it, so §7 must too")
    # …and every retained thing the notice names is on the receipt.
    for what, phrase in _ROW_PHRASE.items():
        assert what in whats, (
            f"the notice's §7 keeps {phrase!r} but the receipt has no {what!r} row — a "
            "receipt that quietly leaves something behind is worse than none")
    assert len(_KEPT) == 6, "six kept rows since 2026-09-04 (erasure record + mail copies joined)"
    # The two rows settled by the founder on 2026-09-04 say what they were settled to say.
    by = {k["what"]: k["why"] for k in _KEPT}
    assert "mobile number" in by["The record that you asked us to erase"], \
        "the erasure record keeps the mobile and SAYS so (hashing was declined)"
    assert "8 years" in by["Email we exchanged"] and "8 years" in by["Tax records for payments made"]
    assert "30 days" in by["Disaster-recovery backups"]
    print("✓ Receipt _KEPT ⇄ notice §7 agree, both directions, six rows")


def test_routes_open_status_seen_export():
    from fastapi.testclient import TestClient
    from api import main as api_main

    c = TestClient(api_main.app)
    # ── OPEN: no X-Aruvi-User at all ──
    r = c.get("/legal/privacy")
    assert r.status_code == 200, "the notice must be readable before any account exists"
    d = r.json()
    assert d["document"]["document_id"] == "privacy_policy"
    assert d["current_version"] == d["document"]["version"]
    assert d["versions"] and d["versions"][-1] == d["current_version"]
    assert c.get("/legal/privacy?version=9.9").status_code == 503, "an unpublished version is refused"
    # ── STATUS before anything is recorded: SILENT (founder: no pop-up for existing
    #    accounts — only a real version bump counts) ──
    uid = "NoticeKumar"
    h = {"X-Aruvi-User": uid}
    s = c.get("/legal/privacy/status", headers=h).json()
    assert s["updated"] is False and s["seen_version"] == ""
    # ── A recorded OLDER version is what raises the bar ──
    acct = api_main.account_repo.load(uid, uid)
    acct.privacy_notice = {"version": "0.0", "seen_at": "2026-01-01T00:00:00+00:00", "context": "test"}
    api_main.account_repo.save(acct)
    s = c.get("/legal/privacy/status", headers=h).json()
    assert s["updated"] is True and s["seen_version"] == "0.0"
    # ── Registration stamps the version current at that moment ──
    assert c.post("/onboarding/verified", headers=h).status_code == 200
    s = c.get("/legal/privacy/status", headers=h).json()
    assert s["updated"] is False and s["seen_version"] == d["current_version"]
    acct = api_main.account_repo.load(uid, uid)
    assert acct.privacy_notice["context"] == "trial_signin" and acct.privacy_notice["seen_at"]
    # ── The seen route re-stamps with its own context ──
    r = c.post("/legal/privacy/seen", json={"context": "updated_note_dismissed"}, headers=h)
    assert r.status_code == 200 and r.json()["status"] == "seen"
    assert api_main.account_repo.load(uid, uid).privacy_notice["context"] == "updated_note_dismissed"
    # ── Rendered in her export, as a fact about the record ──
    from docx import Document
    blob = c.get("/data-rights/export", headers=h).content
    doc = Document(io.BytesIO(blob))
    text = "\n".join(cell.text for t in doc.tables for row in t.rows for cell in row.cells)
    assert "Privacy notice shown" in text and f"Version {d['current_version']}" in text
    # ── Gone with the account ──
    c.post("/data-rights/erase", json={"confirm": "erase", "downloaded_confirmed": True}, headers=h)
    assert c.get("/legal/privacy/status", headers=h).json()["seen_version"] == "", \
        "a fresh account after erasure has no record of a notice shown"
    print("✓ Routes: open GET, status → verified stamps → seen re-stamps → export shows → erase clears")


def test_role_state_city_reach_the_export():
    """3.4 of the considerations: gathered but never rendered was an incomplete access
    right. Now on the account table."""
    from fastapi.testclient import TestClient
    from api import main as api_main
    from docx import Document

    c = TestClient(api_main.app)
    h = {"X-Aruvi-User": "ExportKumar"}
    c.post("/account", json={"name": "Priya Nair", "role": "Teacher", "state": "Tamil Nadu",
                             "city": "Madurai"}, headers=h)
    doc = Document(io.BytesIO(c.get("/data-rights/export", headers=h).content))
    text = "\n".join(cell.text for t in doc.tables for row in t.rows for cell in row.cells)
    assert "Teacher" in text and "Tamil Nadu" in text and "Madurai" in text
    print("✓ Role / state / city are rendered in the export")


if __name__ == "__main__":
    for name, fn in sorted(globals().items()):
        if name.startswith("test_") and callable(fn):
            fn()
    print("\n✅ All privacy-notice tests passed!")
