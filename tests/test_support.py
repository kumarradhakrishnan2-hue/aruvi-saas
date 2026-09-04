"""
Tests for Settings › Support — the email-only support channel (2026-08-27).

Email is the only way a teacher can reach us, so the things worth pinning are the ones
that make a slow channel trustworthy:

  1. THE REFERENCE SERIES. Gapless, never restarts below a number already issued, opens
     at a three-digit offset rather than 1, and lives OUTSIDE any tenant folder — so one
     teacher's erasure cannot hand the next teacher a reference already in use.
  2. THE REQUEST SURVIVES A DEAD MAIL SERVER. Stored first, mailed second; the notifier
     never raises by contract, and a message must never evaporate because SMTP had a bad
     minute.
  3. THE ACKNOWLEDGEMENT SAYS THE THREE THINGS. It arrived · its reference · when a
     human replies. And the window it quotes is the SAME one the screen was told.
  4. HER WORDS ARE HERS. Support messages export with the rest of her data and are
     destroyed by an erase — while the seller's counter is not.
  5. TENANT ISOLATION. One teacher's message never appears in another's listing.

Run standalone:  python3 tests/test_support.py     (also pytest-compatible)
"""
from __future__ import annotations

import io
import os
import sys
import tempfile

sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.abspath(__file__))))

# Throwaway state dir BEFORE importing api.main (see test_account.py). Also forces the
# FileNotifier path — no SMTP credentials, so nothing can leave the machine in a test.
_TMP_STATE = tempfile.mkdtemp(prefix="aruvi-test-state-")
os.environ.setdefault("ARUVI_STATE_DIR", _TMP_STATE)

from aruvi_core.ports import SupportRequest  # noqa: E402
from aruvi_core.adapters.support_repository_file import SupportRepositoryFileImpl  # noqa: E402
from api import mail_templates  # noqa: E402


def _headers(user: str) -> dict:
    return {"X-Aruvi-User": user}


# ── 1. the reference series ────────────────────────────────────────────────────
def test_series_opens_at_the_offset_and_is_gapless():
    with tempfile.TemporaryDirectory() as tmp:
        repo = SupportRepositoryFileImpl(tmp, prefix="ARV-S", start=742)
        first = repo.next_reference()
        assert first == "ARV-S-742", f"the series opens at the offset, got {first}"
        assert [repo.next_reference() for _ in range(3)] == [
            "ARV-S-743", "ARV-S-744", "ARV-S-745"], "gapless, +1 each time"

        # A NEW repository object over the same folder continues the series — the
        # counter is on disk, not in the process.
        assert SupportRepositoryFileImpl(tmp, start=742).next_reference() == "MEY-S-746", "default prefix is MEY-S now; the counter, not the prefix, carries the series"
        print("✓ References open at the three-digit offset and count on, gaplessly")


def test_a_corrupt_counter_restarts_but_never_rewinds():
    """`start` is the FLOOR, not just the seed: the worst a lost counter can do is
    reissue from the offset, never hand out a number below it."""
    with tempfile.TemporaryDirectory() as tmp:
        repo = SupportRepositoryFileImpl(tmp, start=742)
        repo.next_reference()
        with open(os.path.join(tmp, "support", "_series", "support.json"), "w") as f:
            f.write("{ this is not json")
        assert repo.next_reference() == "MEY-S-742", "restarts at the floor, not at 1"
        print("✓ A corrupt counter restarts at the offset and never rewinds below it")


def test_the_counter_is_not_inside_a_tenant_folder():
    """The series belongs to the seller. Inside a teacher's folder it would be destroyed
    by her erasure — the invoice-series reasoning, and the same fix."""
    with tempfile.TemporaryDirectory() as tmp:
        repo = SupportRepositoryFileImpl(tmp, start=742)
        ref = repo.next_reference()
        repo.save(SupportRequest(reference=ref, tenant_id="Kumar1", user_id="Kumar1",
                                 category="problem", message="Hello",
                                 created_at="2026-08-27T09:00:00+00:00"))
        series = os.path.join(tmp, "support", "_series", "support.json")
        assert os.path.exists(series), "the counter has its own home"
        assert not os.path.exists(os.path.join(tmp, "support", "Kumar1", "_series")), \
            "and it is not under any tenant"
        print("✓ The reference counter lives outside every tenant folder")


# ── 2. store and isolation ─────────────────────────────────────────────────────
def test_requests_load_newest_first_and_never_cross_tenants():
    with tempfile.TemporaryDirectory() as tmp:
        repo = SupportRepositoryFileImpl(tmp, start=742)
        for i, (who, when, msg) in enumerate([
                ("Kumar1", "2026-08-25T09:00:00+00:00", "The oldest one"),
                ("Kumar1", "2026-08-27T09:00:00+00:00", "The newest one"),
                ("Priya2", "2026-08-26T09:00:00+00:00", "Priya's private trouble")]):
            repo.save(SupportRequest(reference=repo.next_reference(), tenant_id=who,
                                     user_id=who, category="problem", message=msg,
                                     created_at=when))
        mine = repo.load_all("Kumar1", "Kumar1")
        assert [r.message for r in mine] == ["The newest one", "The oldest one"]
        assert all("Priya" not in r.message for r in mine), "tenant-keyed, isolated"
        assert len(repo.load_all("Priya2", "Priya2")) == 1
        print("✓ Requests come back newest first and never cross tenants")


# ── 3. the acknowledgement ─────────────────────────────────────────────────────
def test_the_acknowledgement_says_the_three_things():
    body = mail_templates.support_acknowledgement(
        name="Latha Menon", reference="ARV-S-742", category="problem",
        message="Chapter 5 shows 19 periods but I asked for 16.",
        reply_days=2, received_on="2026-08-27",
        context={"subject": "social_sciences", "grade": "ix", "chapter": "5"})
    for part in ("subject", "text", "html"):
        assert body[part].strip(), f"the acknowledgement has no {part}"
    # (a) it arrived, (b) its reference — in the subject line too, where a mail client
    #     shows it without her opening anything, (c) when a person replies.
    assert "ARV-S-742" in body["subject"], "the reference is in the subject line"
    assert "reached us" in body["text"]
    assert "2 working days" in body["text"] and "2 working days" in body["html"]
    # Her own words come back to her, verbatim.
    assert "19 periods but I asked for 16" in body["text"]
    assert "19 periods but I asked for 16" in body["html"]
    # And what the app attached on her behalf is SHOWN to her, not just sent to us.
    assert "Social Sciences" in body["text"] and "IX" in body["text"]
    assert "Hello Latha," in body["text"], "greets by first name"
    print("✓ The acknowledgement confirms arrival, names the case, and states the window")


def test_billing_gets_the_firmer_window_and_no_ask_aruvi_line():
    """Ask Aruvi answers "how does this work?", never "why was I charged twice?" — so
    the deflection line is omitted exactly where it would be useless."""
    b = mail_templates.support_acknowledgement(
        name="", reference="ARV-S-743", category="billing", message="Charged twice.",
        reply_days=1)
    assert "1 working day" in b["text"] and "1 working days" not in b["text"]
    assert "Ask Meyy" not in b["text"]
    assert "Hello," in b["text"], "no name is still a greeting, not a blank"
    p = mail_templates.support_acknowledgement(
        name="", reference="ARV-S-744", category="problem", message="Broken.")
    assert "Ask Meyy" in p["text"], "and it IS offered where it can help"
    print("✓ Billing carries the firmer promise; the Ask Meyy line appears only where it helps")


def test_an_unknown_category_is_prettified_not_dropped():
    assert mail_templates.support_category_label("problem") == "Something isn't working"
    assert mail_templates.support_category_label("retired_kind") == "Retired kind"
    assert mail_templates.support_category_label("") == "A message"
    assert mail_templates.reply_window_words(1) == "1 working day"
    assert mail_templates.reply_window_words(2) == "2 working days"
    print("✓ A category we stop offering never erases the case it was filed under")


def test_html_escapes_what_she_typed():
    """The quoted block is her text inside our HTML. Escaped, always — a teacher
    reporting a bug by pasting markup must not be able to rewrite the mail."""
    b = mail_templates.support_acknowledgement(
        name="T", reference="ARV-S-745", category="problem",
        message="<script>alert(1)</script> & \"quotes\"")
    assert "<script>" not in b["html"] and "&lt;script&gt;" in b["html"]
    assert "&amp;" in b["html"]
    print("✓ Her message is escaped into the HTML mail, never injected")


# ── 4. the route ───────────────────────────────────────────────────────────────
def test_route_files_a_case_and_acknowledges_it():
    from fastapi.testclient import TestClient
    from api import main as api_main

    c = TestClient(api_main.app, raise_server_exceptions=False)
    H = _headers("SupportSeeker1")
    # Give her an address, the way checkout or Personal profile would.
    c.post("/account", headers=H, json={"name": "Latha", "email": "latha@example.com"})

    meta = c.get("/support", headers=H).json()
    assert [x["key"] for x in meta["categories"]] == \
        list(mail_templates.SUPPORT_CATEGORIES), "the screen's list, in one order"
    # ★ "Something else" exists, and it is LAST — a list with no escape hatch gets the
    #   nearest wrong bucket picked, and then the category lies to us.
    assert meta["categories"][-1]["key"] == "other"
    assert meta["categories"][-1]["label"] == "Something else"
    assert meta["reply_days"] >= 1 and meta["billing_reply_days"] >= 1
    assert meta["email"] == "latha@example.com", "the screen can say where the reply goes"
    assert meta["requests"] == []

    r = c.post("/support", headers=H, json={
        "category": "plan", "message": "Chapter 5 came back with 19 periods.",
        "context": {"screen": "Settings › Support"}})
    assert r.status_code == 200, r.json()
    got = r.json()
    assert got["reference"].startswith("MEY-S-"), got
    assert got["emailed"] is True, "FileNotifier counts — the message left the app"
    assert got["email"] == "latha@example.com"
    # ★ The window the SCREEN was told and the window the MAIL promises are one value.
    assert got["reply_window"] == mail_templates.reply_window_words(meta["reply_days"])

    stored = api_main.support_repo.load_all("SupportSeeker1", "SupportSeeker1")
    assert len(stored) == 1 and stored[0].reference == got["reference"]
    assert stored[0].acknowledged is True
    assert stored[0].category_label == "Something in a lesson plan looks wrong", \
        "the label she saw is stored, not re-derived later"
    assert stored[0].context.get("screen") == "Settings › Support"

    listed = c.get("/support", headers=H).json()["requests"]
    assert len(listed) == 1 and listed[0]["reference"] == got["reference"]
    print("✓ POST /support files a case, acknowledges it, and the listing shows it")


def test_a_case_is_filed_even_with_no_address_on_the_account():
    """She still gets a reference; the screen then says plainly that we cannot write
    back and offers to fix that — rather than promising a mail that never left."""
    from fastapi.testclient import TestClient
    from api import main as api_main

    c = TestClient(api_main.app, raise_server_exceptions=False)
    H = _headers("SupportNoMail")
    r = c.post("/support", headers=H, json={"category": "problem", "message": "Stuck."})
    assert r.status_code == 200, r.json()
    assert r.json()["emailed"] is False and r.json()["email"] == ""
    # She is given somewhere to write FROM her own mail app instead of a dead end —
    # the common case on trial, where the account holds a mobile and nothing else.
    assert "@" in r.json()["address"], "the support address is offered as the fallback"
    # ★ And it is THE support address (support@meyy.in, founder 2026-09-03) — the same
    # one GET /support tells the screen and the acknowledgement's reply-to — never the
    # sending account, which is the founder's own mailbox.
    from api import config as api_config
    assert r.json()["address"] == api_config.SUPPORT_ADDRESS
    assert c.get("/support", headers=H).json()["address"] == api_config.SUPPORT_ADDRESS
    if not os.environ.get("ARUVI_SUPPORT_ADDRESS"):
        assert api_config.SUPPORT_ADDRESS == "support@meyy.in"
    stored = api_main.support_repo.load_all("SupportNoMail", "SupportNoMail")
    assert len(stored) == 1, "the message is kept regardless"
    assert stored[0].acknowledged is False, "and we know it was never acknowledged"
    print("✓ A teacher with no email on file still gets a real, stored case")


def test_an_empty_or_oversized_message_is_refused_in_her_words():
    from fastapi.testclient import TestClient
    from api import main as api_main

    c = TestClient(api_main.app, raise_server_exceptions=False)
    H = _headers("SupportEdge")
    r = c.post("/support", headers=H, json={"category": "problem", "message": "   "})
    assert r.status_code == 400 and "write your message" in r.json()["detail"]
    r = c.post("/support", headers=H, json={"category": "problem", "message": "x" * 9000})
    assert r.status_code == 400 and "characters" in r.json()["detail"]
    assert api_main.support_repo.load_all("SupportEdge", "SupportEdge") == [], \
        "a refused message files nothing — and burns no reference she was never given"
    # An unknown category falls back rather than 400ing: she has written a real message
    # and a stale client is not her problem.
    r = c.post("/support", headers=H, json={"category": "nonsense", "message": "Hello."})
    assert r.status_code == 200 and r.json()["reference"]
    print("✓ Empty and oversized messages are refused in words she can act on")


def test_support_is_never_gated_on_subscription():
    """★ The one rule that must not drift. A teacher whose SUBSCRIPTION is the broken
    thing has to be able to say so — the same reasoning that keeps data rights ungated
    (§2.5). Asserted with enforcement ON and no entitlement anywhere."""
    from fastapi.testclient import TestClient
    from api import main as api_main

    c = TestClient(api_main.app, raise_server_exceptions=False)
    H = _headers("SupportUnentitled")
    was = api_main.config.ENTITLEMENT_ENFORCED
    api_main.config.ENTITLEMENT_ENFORCED = True
    try:
        r = c.post("/support", headers=H,
                   json={"category": "billing", "message": "I paid and nothing happened."})
        assert r.status_code == 200, r.json()
        # Billing's own, firmer window — resolved server-side, not by the client.
        assert r.json()["reply_days"] == api_main.config.SUPPORT_BILLING_REPLY_DAYS
    finally:
        api_main.config.ENTITLEMENT_ENFORCED = was
    print("✓ Support is reachable with no subscription and enforcement on")


# ── 5. her data rights ─────────────────────────────────────────────────────────
def test_messages_export_with_her_data_and_erase_with_her_account():
    from aruvi_core.adapters.data_rights_service_file import DataRightsServiceFileImpl
    from aruvi_core.ports import Account

    with tempfile.TemporaryDirectory() as tmp:
        svc = DataRightsServiceFileImpl(tmp)
        svc.accounts.save(Account(account_id="Kumar1", tenant_id="Kumar1",
                                  display_name="Kumar", email="k@example.com"))
        ref = svc.support.next_reference()
        svc.support.save(SupportRequest(
            reference=ref, tenant_id="Kumar1", user_id="Kumar1", category="plan",
            category_label="Something in a lesson plan looks wrong",
            message="The shipwreck cards are missing.",
            created_at="2026-08-27T09:00:00+00:00"))

        text = _docx_text(svc.export("Kumar1", "Kumar1"))
        assert ref in text, "her reference is in her export"
        assert "shipwreck cards are missing" in text, "and so are her own words"
        assert "Something in a lesson plan looks wrong" in text, "with the label she saw"

        receipt = svc.erase("Kumar1", "Kumar1")
        assert "support messages" in receipt.erased, receipt.erased
        assert svc.support.load_all("Kumar1", "Kumar1") == [], "gone"
        # ★ But the SELLER's counter is not — the next reference follows the last one
        #   issued rather than repeating it.
        assert svc.support.next_reference() != ref, \
            "an erase must never make the series hand out a used reference"
        print("✓ Support messages export with her data, erase with her account, "
              "and never take the counter with them")


def _docx_text(blob: bytes) -> str:
    """All text in the document — paragraphs AND table cells (same helper as
    test_data_rights; the export carries content in both)."""
    from docx import Document
    doc = Document(io.BytesIO(blob))
    parts = [p.text for p in doc.paragraphs]
    for tbl in doc.tables:
        for row in tbl.rows:
            for cell in row.cells:
                parts.append(cell.text)
    return "\n".join(parts)


if __name__ == "__main__":
    test_series_opens_at_the_offset_and_is_gapless()
    test_a_corrupt_counter_restarts_but_never_rewinds()
    test_the_counter_is_not_inside_a_tenant_folder()
    test_requests_load_newest_first_and_never_cross_tenants()
    test_the_acknowledgement_says_the_three_things()
    test_billing_gets_the_firmer_window_and_no_ask_aruvi_line()
    test_an_unknown_category_is_prettified_not_dropped()
    test_html_escapes_what_she_typed()
    test_route_files_a_case_and_acknowledges_it()
    test_a_case_is_filed_even_with_no_address_on_the_account()
    test_an_empty_or_oversized_message_is_refused_in_her_words()
    test_support_is_never_gated_on_subscription()
    test_messages_export_with_her_data_and_erase_with_her_account()
    print("\n✅ All support tests passed!")
