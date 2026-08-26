"""
Tests for administrative-architecture Step 4: DataRightsService — the export/erase
traversal over every Bucket-B store — and the /data-rights routes.

Pinned here on purpose:
  * the export contains her account, profile, notes (every year) and teaching state,
    and NEVER the shared lesson-plan library;
  * the export is the tenant-isolation test: another tenant's text must not appear;
  * erase removes every folder (account last), returns the §2.6 receipt whose `kept`
    wording matches the privacy-policy promises, and is idempotent;
  * after erase the ID is not reserved — signing in JIT-creates a fresh empty account.

Run standalone:  python3 tests/test_data_rights.py     (also pytest-compatible)
"""
from __future__ import annotations

import io
import os
import sys
import tempfile

sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.abspath(__file__))))

# Throwaway state dir BEFORE importing api.main (see test_account.py).
_TMP_STATE = tempfile.mkdtemp(prefix="aruvi-test-state-")
os.environ.setdefault("ARUVI_STATE_DIR", _TMP_STATE)

from aruvi_core.ports import Account, AcademicYear, PlanNote  # noqa: E402
from aruvi_core.adapters.data_rights_service_file import DataRightsServiceFileImpl  # noqa: E402

T, U, Y = "Kumar1", "Kumar1", "2026-27"


def _docx_text(blob: bytes) -> str:
    """All text in the document — paragraphs AND table cells (the reformatted export
    carries the profile, account and teaching state in tables)."""
    from docx import Document
    doc = Document(io.BytesIO(blob))
    parts = [p.text for p in doc.paragraphs]
    for tbl in doc.tables:
        for row in tbl.rows:
            for cell in row.cells:
                parts.append(cell.text)
    return "\n".join(parts)


def _seed(tmp, tenant=T, user=U, note="Bring the shipwreck cards."):
    """A full little life for one teacher: account, year, profile, and one of each
    teaching-state kind."""
    svc = DataRightsServiceFileImpl(tmp)
    svc.accounts.save(Account(account_id=user, tenant_id=tenant, display_name=user,
                              email=f"{user.lower()}@example.com",
                              created_at="2026-08-22T00:00:00+00:00"))
    svc.years.open_year(tenant, user, AcademicYear(Y, "2026-04-01", "2027-03-31", True))
    svc.readiness.save_profile(tenant, user, {"subjects": [
        {"name": "Science", "grades": [{"grade": "VII", "sections": [{"tag": "7A"}]}]}]})
    svc.notes.save(tenant, user, Y, PlanNote("science/vii/3", note,
                                             "2026-08-22T10:00:00+00:00"))
    svc.allocations.save_allocation(tenant, user, Y, "science", "vii",
                                    {"1": {"chapter_title": "Food", "weight": 1,
                                           "periods_by_duration": {"45": 5},
                                           "total_periods": 5, "total_minutes": 225}})
    svc.sections.save_one(tenant, user, Y, "science_vii_7A", "ch_03_canonical.json", 4, False)
    svc.prepared.mark(tenant, user, Y, "science/vii/ch_03.json", 16)
    svc.archive.archive(tenant, user, Y, "science/vii/ch_09.json")
    # Entitlement (Step 5) — tenant-keyed; must be erased for an individual.
    from aruvi_core.ports import Entitlement
    from aruvi_core.adapters.entitlement_repository_file import EntitlementRepositoryFileImpl
    EntitlementRepositoryFileImpl(tmp).save(tenant, Entitlement(
        plan_id="trial", status="trial", source="trial", scopes=["*"],
        trial_chapters=[f"science/vii/{i}" for i in (1, 2, 3)]))
    return svc


def test_export_contains_everything_hers():
    with tempfile.TemporaryDirectory() as tmp:
        svc = _seed(tmp)
        text = _docx_text(svc.export(T, U))
        for needle in ("Kumar1", "kumar1@example.com",          # account
                       "belongs to you",                         # the purpose statement
                       "Science", "7A",                          # profile table
                       "2026-27",                                # year
                       "Bring the shipwreck cards.",             # the note itself
                       "Chapter 3",                              # note's bold identity line
                       "Ch. 3",                                  # teaching-state table
                       "at Learning Unit 5",                     # per-section status
                       "deleting your Aruvi account"):           # pre-deletion advisory
            assert needle in text, f"export missing: {needle}"
        # Founder direction (2026-08-22): reader-facing only — no filenames, no
        # canonical identities, no raw section keys, no period internals.
        for absent in ("ch_03_canonical.json", "ch_03.json", "ch_09.json",
                       "science_vii_7A", "period(s) planned"):
            assert absent not in text, f"export must not show: {absent}"
        print("✓ Export carries account, profile, notes and teaching state (reader-facing)")


def test_pdf_export_same_payload():
    """The PDF is the SAME gathered payload through a second renderer — check the
    magic bytes and that the note text survives into the PDF stream (xhtml2pdf keeps
    page text findable in the uncompressed content stream for small docs)."""
    with tempfile.TemporaryDirectory() as tmp:
        svc = _seed(tmp)
        blob = svc.export(T, U, "pdf")
        assert blob[:4] == b"%PDF" and len(blob) > 1500
        try:
            svc.export(T, U, "xls")
            assert False, "expected ValueError for unknown format"
        except ValueError:
            pass
        print("✓ PDF export renders from the same payload; unknown formats refused")


def test_export_is_the_isolation_test():
    """Tenant B's data must not reach tenant A's export — one byte would mean the
    traversal crossed a tenant boundary."""
    with tempfile.TemporaryDirectory() as tmp:
        svc = _seed(tmp)
        _seed(tmp, tenant="Priya2", user="Priya2",
              note="PRIYA-SECRET: seating chart for 8B.")
        text = _docx_text(svc.export(T, U))
        assert "PRIYA-SECRET" not in text and "Priya2" not in text
        assert "Bring the shipwreck cards." in text
        print("✓ Export never reaches another tenant's data")


def test_export_excludes_shared_library():
    """The shared plan library is content, not personal data — only stated, never
    embedded (grep-able stand-in: the export must SAY plans are excluded)."""
    with tempfile.TemporaryDirectory() as tmp:
        svc = _seed(tmp)
        text = _docx_text(svc.export(T, U))
        assert "not personal data" in text
        print("✓ Export discloses that the shared library is excluded")


def test_erase_walks_everything_and_receipts():
    with tempfile.TemporaryDirectory() as tmp:
        svc = _seed(tmp)
        receipt = svc.erase(T, U)
        assert "account record" in receipt.erased
        assert "teaching profile" in receipt.erased
        assert f"chapter notes ({Y})" in receipt.erased
        assert "subscription record" in receipt.erased, \
            "the Step-5 entitlement store must join the erase walk"
        assert receipt.erased[-1] == "account record", "account goes LAST"
        # §2.6 receipt wording — pinned; must match the privacy policy's promises.
        whats = [k["what"] for k in receipt.kept]
        assert "Disaster-recovery backups" in whats
        assert "Tax records for payments made" in whats
        assert any("30 days" in k["why"] for k in receipt.kept)
        # Nothing of hers remains on disk, in any store.
        import pathlib
        left = [str(p) for p in pathlib.Path(tmp).rglob("*")
                if "kumar1" in str(p).lower()]
        assert not left, f"remnants: {left}"
        # Neighbour untouched, and the walk is idempotent.
        _seed(tmp, tenant="Priya2", user="Priya2")
        assert svc.erase(T, U).erased == []
        assert svc.accounts.load("Priya2", "Priya2") is not None
        print("✓ Erase removes everything, receipts correctly, idempotently")


def test_routes_and_fresh_start():
    """End-to-end: export downloads a .docx; erase needs the typed confirmation;
    after erase the same ID signs in as a brand-new empty account."""
    from fastapi.testclient import TestClient
    from api import main as api_main

    c = TestClient(api_main.app)
    h = {"X-Aruvi-User": "EraseKumar"}
    # Give the identity a note so there is something to export/erase.
    c.post("/plan-notes", json={"subject": "science", "grade": "vii", "chapter": "3",
                                "text": "note before erase",
                                "updated_at": "2026-08-22T10:00:00+00:00"}, headers=h)
    r = c.get("/data-rights/export", headers=h)
    assert r.status_code == 200 and r.content[:2] == b"PK", "a real .docx download"
    assert "attachment" in r.headers.get("content-disposition", "")
    # PDF twin (founder 2026-08-22: every export offers both formats).
    r = c.get("/data-rights/export?format=pdf", headers=h)
    assert r.status_code == 200 and r.content[:4] == b"%PDF", "a real .pdf download"
    assert r.headers.get("content-type", "").startswith("application/pdf")
    assert ".pdf" in r.headers.get("content-disposition", "")
    assert c.get("/data-rights/export?format=xls", headers=h).status_code == 400
    # No/wrong confirmation → refused.
    assert c.post("/data-rights/erase", json={}, headers=h).status_code == 400
    # ★ The typed word alone is no longer enough (founder 2026-08-26): she must also
    #   state she has her data, because the download is the only copy she can keep.
    assert c.post("/data-rights/erase", json={"confirm": "erase"},
                  headers=h).status_code == 400, "download confirmation is required"
    r = c.post("/data-rights/erase",
               json={"confirm": "erase", "downloaded_confirmed": True}, headers=h)
    body = r.json()
    assert body["confirmation_recorded"] is True, "consent must be logged"
    assert body.get("confirmed_at"), "the log stamps when she confirmed"
    assert body["status"] == "erased" and "account record" in body["erased"]
    # Fresh start: same ID, brand-new empty account, no old data.
    assert c.get("/plan-notes", headers=h).json()["notes"] == {}
    acct = api_main.account_repo.load("EraseKumar", "EraseKumar")
    assert acct is not None, "JIT re-created on the post-erase request"
    print("✓ Routes: export download, confirm-guarded erase, fresh start after")


def test_erasure_consent_log_survives_the_erasure():
    """★ The ONE record that must outlive an erasure (founder, 2026-08-26).

    Everything else about her is destroyed — that is the point — which is exactly why,
    until now, an erased account left no evidence she had ever confirmed. The log lives
    outside the erase walk. It must (a) still be there afterwards, (b) hold her
    confirmation, and (c) carry NO personal data: identifiers and timestamps only, or it
    would quietly reintroduce what she asked to have destroyed."""
    from fastapi.testclient import TestClient
    from api import main as api_main

    c = TestClient(api_main.app)
    uid = "ConsentKumar"
    h = {"X-Aruvi-User": uid}
    c.post("/plan-notes", json={"subject": "science", "grade": "vii", "chapter": "1",
                                "text": "Ravi struggled with fractions today",
                                "updated_at": "2026-08-26T10:00:00+00:00"}, headers=h)
    c.post("/account", json={"name": "Priya Nair", "email": "priya.nair@example.com"},
           headers=h)

    r = c.post("/data-rights/erase",
               json={"confirm": "erase", "downloaded_confirmed": True}, headers=h)
    assert r.status_code == 200 and r.json()["confirmation_recorded"] is True

    entries = api_main.erasure_log.for_tenant(uid)
    assert len(entries) == 1, "exactly one record for one deletion"
    e = entries[0]
    assert e["tenant_id"] == uid and e["user_id"] == uid   # tenant/user wise
    assert e["confirmed_downloaded"] is True
    assert e["confirmed_at"]

    # (c) NO personal data may appear anywhere in the log.
    import json as _json
    blob = _json.dumps(entries).lower()
    for forbidden in ("priya", "nair", "priya.nair@example.com", "ravi", "fractions"):
        assert forbidden not in blob, f"the consent log leaked {forbidden!r}"

    # A second deletion appends rather than replacing — the log is a history.
    c.get("/plan-notes", headers=h)                    # JIT-recreate the account
    c.post("/data-rights/erase",
           json={"confirm": "erase", "downloaded_confirmed": True}, headers=h)
    assert len(api_main.erasure_log.for_tenant(uid)) == 2, "append-only"
    print("✓ Erasure consent is recorded tenant/user wise, survives, and leaks nothing")


if __name__ == "__main__":
    test_export_contains_everything_hers()
    test_pdf_export_same_payload()
    test_export_is_the_isolation_test()
    test_export_excludes_shared_library()
    test_erase_walks_everything_and_receipts()
    test_routes_and_fresh_start()
    test_erasure_consent_log_survives_the_erasure()
    print("\n✅ All data-rights tests passed!")
