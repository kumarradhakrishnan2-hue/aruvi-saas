"""
Tests for the Year Plan Word export — aruvi_core/export_year_plan_docx.py.

What is pinned here, and why each one is worth a test:

  * THE DOCUMENT SAYS WHAT THE SCREEN SAID. The whole design of this export is that
    the client posts its own model and the server renders it — so the one thing that
    must never regress is the server quietly substituting a number of its own. The
    totals test asserts the payload's totals appear even when they disagree with the
    column beneath them.
  * A MISSING VALUE IS AN EM-DASH, NEVER A ZERO. The Support `metaErr` rule applied
    to a document: it may say it does not know, it may never invent an answer about
    her record. A chapter with no suggestion and a chapter suggested 0 periods are
    different facts and must read differently.
  * "SET" SURVIVES. A legacy prepare stored no `prepared_periods`; the screen shows
    "set" for it, and a dash there would misreport a prepared chapter as untouched.
  * THE TABLE IS THE WHOLE DOCUMENT. No competencies, no effort-index values, no
    summary strip — if it is not on the pane it is not in the file.

Run standalone:  python3 tests/test_year_plan_export.py     (also pytest-compatible)
"""
from __future__ import annotations

import io
import os
import re
import sys

sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.abspath(__file__))))

from aruvi_core.export_year_plan_docx import export_year_plan_docx  # noqa: E402

try:
    from docx import Document
except ImportError:  # pragma: no cover - the module itself would not import either
    Document = None


PAYLOAD = {
    "subject": "Social Sciences",
    "grade": "ix",
    "budget": 245,
    "generated_at": "2026-08-30T09:15:00",
    "rows": [
        {"n": 1, "title": "The Land of Rivers", "sug": 21, "plan": 20,
         "prepared": True, "awaited": False},
        # Prepared, but a legacy prepare with no periods recorded → "set".
        {"n": 2, "title": "Early Societies", "sug": 18, "plan": None,
         "prepared": True, "awaited": False},
        # Not prepared → dash on the plan side.
        {"n": 3, "title": "Maps and Scale", "sug": 14, "plan": None,
         "prepared": False, "awaited": False},
        # Budgeted but unpublished — carries periods, can carry no plan.
        {"n": 4, "title": "Book awaited", "sug": 12, "plan": None,
         "prepared": False, "awaited": True},
        # No suggestion available at all (no weight, no budget) → dash, not 0.
        {"n": 5, "title": "Living Together", "sug": None, "plan": 9,
         "prepared": True, "awaited": False},
        # A real zero — a chapter the distribution gave nothing to. NOT a dash.
        {"n": 6, "title": "Revision", "sug": 0, "plan": 0,
         "prepared": True, "awaited": False},
    ],
    "sug_total": 65,
    "plan_total": 29,
}


def _doc(payload=None):
    data = export_year_plan_docx(payload if payload is not None else PAYLOAD)
    assert isinstance(data, bytes) and len(data) > 1000, "docx bytes look empty"
    return Document(io.BytesIO(data))


def _table_rows(doc):
    """The year-plan table is the one with 4 columns (the header table has 2)."""
    for t in doc.tables:
        if len(t.columns) == 4:
            return [[c.text.strip() for c in r.cells] for r in t.rows]
    raise AssertionError("no 4-column table found in the document")


def _all_text(doc):
    parts = [p.text for p in doc.paragraphs]
    for t in doc.tables:
        for r in t.rows:
            parts.extend(c.text for c in r.cells)
    return "\n".join(parts)


def test_opens_as_a_word_document():
    doc = _doc()
    assert doc.tables, "expected at least the header table and the year-plan table"


def test_header_identifies_the_pane():
    text = _all_text(_doc())
    assert "Aruvi" in text
    assert "Year plan" in text
    # Class, not Grade — the teacher-facing word on this pane.
    assert "Class IX" in text
    assert "Social Sciences" in text
    assert "30-August-2026" in text


def test_columns_are_the_screens_three():
    rows = _table_rows(_doc())
    head = [c.upper() for c in rows[0]]
    assert head == ["#", "CHAPTER", "SUGGESTED PERIODS", "YOUR PLAN"]


def test_every_chapter_is_present_in_order():
    rows = _table_rows(_doc())
    titles = [r[1] for r in rows[1:-1]]
    assert titles == [r["title"] for r in PAYLOAD["rows"]]
    assert rows[1][0] == "01", "chapter numbers are zero-padded like the pane"


def test_suggested_column_is_the_payloads():
    rows = _table_rows(_doc())
    assert [r[2] for r in rows[1:-1]] == ["21", "18", "14", "12", "—", "0"]


def test_missing_is_a_dash_and_zero_is_a_zero():
    """The one that matters: a chapter with no suggestion and a chapter suggested
    ZERO periods are different facts. Conflating them invents an answer."""
    rows = _table_rows(_doc())
    by_title = {r[1]: r for r in rows[1:-1]}
    assert by_title["Living Together"][2] == "—"
    assert by_title["Revision"][2] == "0"


def test_plan_column_mirrors_the_panes_four_way_render():
    rows = _table_rows(_doc())
    by_title = {r[1]: r for r in rows[1:-1]}
    assert by_title["The Land of Rivers"][3] == "20"   # periods recorded
    assert by_title["Early Societies"][3] == "set"     # prepared, no periods (legacy)
    assert by_title["Maps and Scale"][3] == "—"        # not prepared
    assert by_title["Book awaited"][3] == "—"          # can carry no plan
    assert by_title["Revision"][3] == "0"              # a real zero commitment


def test_total_row_carries_the_payloads_totals_not_a_resum():
    """Deliberately sabotaged: the payload's totals do NOT equal the columns above
    them. The document must still say what the pane said — the server never
    substitutes arithmetic of its own."""
    p = dict(PAYLOAD, sug_total=999, plan_total=111)
    rows = _table_rows(_doc(p))
    tot = rows[-1]
    assert tot[1] == "Total periods"
    assert tot[2] == "999"
    assert tot[3] == "111"


def test_note_explains_the_two_columns_and_names_the_budget():
    text = _all_text(_doc())
    assert "245 periods" in text
    assert "all 6 chapters" in text
    assert "Suggested periods" in text and "Your plan" in text


def test_note_matches_the_panes_own_words():
    """★ THE CROSS-LANGUAGE GUARD. The note is a second copy of `.yp-note` in
    YearPlan.jsx, and the first build silently dropped its closing sentence ("To know
    how Aruvi suggests, refer to Ask Aruvi time allocation section.") — a whole line of
    the teacher's explanation missing from the file, with nothing to catch it.

    This reads the JSX, cuts out the interpolated bits, and asserts every continuous RUN
    of the pane's fixed prose survives into the document. Cutting at the interpolations
    rather than at sentence boundaries is what makes it robust: the budget clause and the
    chapter count legitimately differ between the two, so the note is compared as the
    fixed spans BETWEEN them — a brittle test is a deleted test.
    """
    jsx_path = os.path.join(os.path.dirname(os.path.dirname(os.path.abspath(__file__))),
                            "web", "app", "components", "YearPlan.jsx")
    if not os.path.exists(jsx_path):  # engine-only checkout
        print("  skip note-parity — web/ not present")
        return
    src = open(jsx_path, encoding="utf-8").read()
    m = re.search(r'<p className="yp-note">(.*?)</p>', src, re.S)
    assert m, "could not find the .yp-note paragraph in YearPlan.jsx"

    # Replace every {…} expression — NESTED ones included, which a regex cannot do — with a
    # sentinel, so the fixed prose either side of it stays intact and separable.
    SENT = "\x00"
    out, depth = [], 0
    for ch in m.group(1):
        if ch == "{":
            if depth == 0:
                out.append(SENT)
            depth += 1
        elif ch == "}":
            depth = max(0, depth - 1)
        elif depth == 0:
            out.append(ch)
    note = "".join(out)
    note = re.sub(r"<[^>]+>", " ", note)                 # <b>, <br />, fragments
    note = note.replace("&rsquo;", "'").replace("&mdash;", "—")
    note = re.sub(r"[ \t\r\n]+", " ", note)
    note = re.sub(r"\s+([,.;])", r"\1", note)            # tag removal leaves " ,"

    doc_text = re.sub(r"\s+", " ", _all_text(_doc()))
    missing = []
    for run in note.split(SENT):
        r = run.strip()
        # Anything shorter than a clause is a connector left beside an interpolation
        # ("spread across all"), not a claim worth pinning.
        if len(r) < 25:
            continue
        if r not in doc_text:
            missing.append(r)
    assert not missing, "the export is missing prose the pane shows: " + repr(missing)


def test_note_carries_the_ask_aruvi_pointer():
    """Pinned by name because this is the sentence that went missing."""
    assert ("To know how Aruvi suggests, refer to Ask Aruvi time allocation section."
            in _all_text(_doc()))


def test_no_budget_reads_as_your_periods_not_a_number():
    p = dict(PAYLOAD, budget=None)
    text = _all_text(_doc(p))
    assert "your periods" in text
    assert "budget of" not in text


def test_the_table_is_the_whole_document():
    """No competencies, no effort-index values, no summary strip — the export is the
    pane, and the pane deliberately shows none of these."""
    text = _all_text(_doc()).lower()
    for forbidden in ("competency", "competencies", "effort index", "executive summary",
                      "weight", "period types"):
        assert forbidden not in text, f"{forbidden!r} leaked into the year-plan export"


def test_empty_year_still_renders():
    """A subject with no chapters must not throw — the pane's own empty state stops
    this from being reachable today, but a document generator that raises on an empty
    list is one refactor away from a 500."""
    p = {"subject": "science", "grade": "vii", "budget": None, "rows": [],
         "sug_total": 0, "plan_total": 0}
    rows = _table_rows(_doc(p))
    assert rows[-1][1] == "Total periods"
    assert len(rows) == 2, "header + total only"


def test_subject_slug_is_displayed_readably():
    p = dict(PAYLOAD, subject="social_sciences", grade="VII")
    text = _all_text(_doc(p))
    assert "Social Sciences" in text
    assert "Class VII" in text


def _main():
    tests = [(n, f) for n, f in sorted(globals().items())
             if n.startswith("test_") and callable(f)]
    failed = 0
    for name, fn in tests:
        try:
            fn()
            print(f"  ok   {name}")
        except AssertionError as e:
            failed += 1
            print(f"  FAIL {name}: {e}")
    print(f"\n{len(tests) - failed}/{len(tests)} passed")
    return 1 if failed else 0


if __name__ == "__main__":
    sys.exit(_main())
